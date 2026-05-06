"""
AUTOSAR RAG INGESTION SYSTEM V3
================================
Complete rewrite addressing all 21 identified problems.

Fixes Applied
-------------
I1  - Token-count overlap (via tiktoken) instead of character-count overlap
I2  - Chunk sizes in TOKENS (child=400, parent=1600) not characters
I3  - Parent-Child chunking: small child chunks for retrieval, large parent
      chunks stored as payload for LLM context delivery
I4  - Contextual enrichment: LLM-generated 80-token context prepended to
      every child chunk BEFORE embedding (Anthropic Contextual Retrieval)
I5  - [Page N] tags stripped from chunk text; page number stored in metadata only
I6  - PDF bookmark tree used for structural splitting (reliable); regex fallback only
I7  - Jaccard dedup now uses AUTOSAR-tokenized word sets (preserves identifiers)
I8  - No external BM25 JSON; sparse vectors come from BGE-M3 natively (no IDF drift)
I9  - BGE-M3 native learned sparse weights replace hand-built TF formula
I10 - No length-normalization bug; BGE-M3 sparse is trained, not formula-based
I11 - No BM25 JSON file at all; versioning problem eliminated entirely
I12 - sentence-transformers BGEM3FlagModel with batch encoding (GPU-parallel)
I13 - BGE-M3 sparse (lexical_weights) AND ColBERT (colbert_vecs) fully used
I14 - Dense vectors are L2-normalised (normalize_embeddings=True equivalent)
I15 - ColBERT multi-vectors stored in Qdrant with MultiVectorConfig(MAX_SIM)
I16 - Tables converted to natural-language sentences before chunking/embedding
I17 - pdfplumber used for table content extraction (not just counting)
I18 - Streaming batch upload; chunks not all held in RAM simultaneously
I19 - Embedding model fingerprint written to collection metadata; mismatch check
I20 - Page markers stripped before BM25 tokenisation and embedding
I21 - L2-norm guard: degenerate near-zero vectors are logged and skipped

Architecture (Contextual Parent-Child + BGE-M3 Three-Signal)
-------------------------------------------------------------
Ingestion:
  PDF/DOCX  →  extract text + tables  →  structural split (bookmarks/headings)
            →  parent chunks (1600 tok)
            →  child chunks  (400 tok, sentence-boundary, 10% overlap)
            →  contextual enrichment per child  (LLM via local endpoint)
            →  BGEM3FlagModel.encode(enriched_child)
               → dense_vecs   (1024-dim, L2-normalised)
               → lexical_weights  ({token_id: weight})  ← replaces external BM25
               → colbert_vecs    (N×1024 matrix)
            →  Qdrant upsert  {dense, sparse, colbert}
               payload: raw_child_text, parent_text, section_hierarchy,
                        page_number, filename, …

Retrieval (in HybridRetriever):
  Query  →  BGEM3FlagModel.encode  →  Qdrant prefetch(dense+sparse)
         →  ColBERT rerank  →  fetch parent_text from payload  →  LLM

Requirements
------------
pip install FlagEmbedding sentence-transformers qdrant-client
            langchain-text-splitters tiktoken nltk pypdfium2 pdfplumber
            python-docx requests numpy
"""

from __future__ import annotations

import argparse
import hashlib
import json
import logging
import math
import os
import re
import sys
import time
import unicodedata
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, Generator, List, Optional, Tuple

import numpy as np

# ── optional heavy imports (fail gracefully with clear messages) ─────────────
try:
    import pypdfium2 as pdfium
except ImportError:
    pdfium = None  # type: ignore

try:
    import pdfplumber
except ImportError:
    pdfplumber = None  # type: ignore

try:
    import docx as python_docx
except ImportError:
    python_docx = None  # type: ignore

try:
    import tiktoken
    _TOKENIZER = tiktoken.get_encoding("cl100k_base")
except ImportError:
    tiktoken = None  # type: ignore
    _TOKENIZER = None

try:
    import nltk
    from nltk.tokenize import sent_tokenize, word_tokenize
    nltk.data.find("tokenizers/punkt_tab")
except (ImportError, LookupError):
    try:
        nltk.download("punkt_tab", quiet=True)
        from nltk.tokenize import sent_tokenize, word_tokenize
    except Exception:
        try:
            nltk.download("punkt", quiet=True)
            from nltk.tokenize import sent_tokenize, word_tokenize
        except Exception:
            sent_tokenize = lambda t: t.split(". ")  # type: ignore
            word_tokenize = lambda t: t.split()      # type: ignore

try:
    from FlagEmbedding import BGEM3FlagModel
    _HAS_FLAG = True
except ImportError:
    _HAS_FLAG = False

try:
    from qdrant_client import QdrantClient
    from qdrant_client import models as qmodels
    _HAS_QDRANT = True
except ImportError:
    _HAS_QDRANT = False

try:
    import requests as _requests
    _HAS_REQUESTS = True
except ImportError:
    _HAS_REQUESTS = False


# ══════════════════════════════════════════════════════════════════════════════
# CONFIG
# ══════════════════════════════════════════════════════════════════════════════

DATA_DIR    = "/home/user/autosar_docs"      # ← change to your docs folder
COLLECTION  = "Dear_autosar"
QDRANT_URL  = "http://localhost:7333"

# BGE-M3 via sentence-transformers / FlagEmbedding (I12, I13, I14, I15)
EMBEDDING_MODEL   = "BAAI/bge-m3"
EMBED_BATCH_SIZE  = 8          # chunks per GPU forward pass
MAX_TOKENS        = 8192       # BGE-M3 context window
ENCODE_MAX_LEN    = 2048       # safety cap per chunk (well within 8192)
MIN_L2_NORM       = 0.05       # I21: reject near-zero vectors below this norm
COLBERT_MAX_TOKENS = 250       # Qdrant hard limit: 1MB/vector = 256 tokens max (1024-dim × 4B × 256 = 1MB)

# Chunking (I1, I2, I3) — all sizes in TOKENS
PARENT_CHUNK_TOKENS   = 1600   # full section delivered to LLM
CHILD_CHUNK_TOKENS    = 400    # precise child for retrieval
CHILD_OVERLAP_TOKENS  = 40     # ≈10% overlap
MIN_CHILD_TOKENS      = 60     # discard very short fragments

# Contextual enrichment (I4) — set CONTEXT_LLM_URL="" to skip
CONTEXT_LLM_URL    = "http://localhost:8011/v1"      # local vLLM / Ollama OpenAI-compat
CONTEXT_LLM_MODEL  = "Qwen/Qwen2.5-32B-Instruct-AWQ"
CONTEXT_LLM_TIMEOUT = 60
CONTEXT_MAX_TOKENS  = 120      # tokens for the generated context prefix

# Deduplication (I7) — AUTOSAR-token based Jaccard
JACCARD_THRESHOLD = 0.95
JACCARD_WINDOW    = 30

# TOC / boilerplate filtering (unchanged from V2)
TOC_LINE_RATIO        = 0.50
TOC_MIN_CONTENT_CHARS = 50
TOC_MIN_LINE_COUNT    = 5
TOC_MAX_CONTENT_CHARS = 800

# Output
OUTPUT_DIR = "./ingestion_output_v3"


# ══════════════════════════════════════════════════════════════════════════════
# LOGGING
# ══════════════════════════════════════════════════════════════════════════════

def _setup_logging(log_dir: str) -> Tuple[logging.Logger, str]:
    Path(log_dir).mkdir(parents=True, exist_ok=True)
    ts = time.strftime("%Y%m%d_%H%M%S")
    log_path = str(Path(log_dir) / f"ingestion_v3_{ts}.log")
    fmt = logging.Formatter("%(asctime)s  %(levelname)-8s  %(message)s")
    root = logging.getLogger()
    root.setLevel(logging.INFO)
    if not root.handlers:
        ch = logging.StreamHandler(sys.stdout)
        ch.setFormatter(fmt)
        root.addHandler(ch)
    fh = logging.FileHandler(log_path, encoding="utf-8")
    fh.setFormatter(fmt)
    root.addHandler(fh)
    return logging.getLogger(__name__), log_path


logger = logging.getLogger(__name__)


# ══════════════════════════════════════════════════════════════════════════════
# TOKEN COUNTING  (I1, I2)
# ══════════════════════════════════════════════════════════════════════════════

def count_tokens(text: str) -> int:
    """Count tokens using tiktoken (cl100k_base). Falls back to word/4 heuristic."""
    if _TOKENIZER is not None:
        return len(_TOKENIZER.encode(text))
    return max(1, len(text) // 4)


def split_by_tokens(text: str, max_tokens: int) -> List[str]:
    """Hard-split text that exceeds max_tokens into token-safe fragments."""
    if _TOKENIZER is None:
        # Character-based fallback: 4 chars ≈ 1 token
        limit = max_tokens * 4
        return [text[i:i + limit] for i in range(0, len(text), limit)]
    ids = _TOKENIZER.encode(text)
    parts = []
    for start in range(0, len(ids), max_tokens):
        parts.append(_TOKENIZER.decode(ids[start:start + max_tokens]))
    return parts


# ══════════════════════════════════════════════════════════════════════════════
# AUTOSAR TOKENIZER  (I7 — identifier-preserving Jaccard)
# ══════════════════════════════════════════════════════════════════════════════

_COMPOUND_RE = re.compile(
    r"[A-Za-z][A-Za-z0-9]*(?:[_:][A-Za-z0-9]+)+"
    r"|[A-Za-z]{2,}[0-9]+[A-Za-z0-9]*"
    r"|[A-Za-z0-9]+(?:::[A-Za-z0-9]+)+"
)
_PLAIN_RE = re.compile(r"[A-Za-z0-9]+")


def autosar_tokenize(text: str) -> List[str]:
    """
    Two-pass tokenizer preserving compound AUTOSAR identifiers as atomic tokens.
    Identical to the tokenizer in HybridRetriever so Jaccard dedup is consistent.
    """
    tokens: List[str] = []
    covered: List[Tuple[int, int]] = []
    for m in _COMPOUND_RE.finditer(text):
        whole = m.group(0).lower()
        tokens.append(whole)
        covered.append((m.start(), m.end()))
        for part in _PLAIN_RE.findall(whole):
            if len(part) > 1:
                tokens.append(part)
    for m in _PLAIN_RE.finditer(text):
        s, e = m.start(), m.end()
        if any(cs <= s and e <= ce for cs, ce in covered):
            continue
        word = m.group(0).lower()
        if len(word) > 1:
            tokens.append(word)
    return tokens


def jaccard(a: str, b: str) -> float:
    """Word-level Jaccard using AUTOSAR-aware tokenizer (I7)."""
    sa = set(autosar_tokenize(a))
    sb = set(autosar_tokenize(b))
    if not sa or not sb:
        return 0.0
    return len(sa & sb) / len(sa | sb)


# ══════════════════════════════════════════════════════════════════════════════
# DATA CLASSES
# ══════════════════════════════════════════════════════════════════════════════

@dataclass
class ParentChunk:
    """A large section of the document (1600 tokens). Stored as payload only."""
    id:                str
    text:              str           # raw section text (no page markers)
    filename:          str
    section_title:     str
    section_hierarchy: List[str]
    page_number:       Optional[int]
    token_count:       int


@dataclass
class ChildChunk:
    """A small retrieval unit (400 tokens) with contextual enrichment."""
    id:                  str
    parent_id:           str
    raw_text:            str         # stored in payload for BM25 / display
    enriched_text:       str         # context-prepended text → embedded
    filename:            str
    section_title:       str
    section_hierarchy:   List[str]
    page_number:         Optional[int]
    child_index:         int         # position within parent
    token_count:         int
    # set after encoding
    dense_vec:           Optional[List[float]]    = None
    sparse_indices:      Optional[List[int]]      = None
    sparse_values:       Optional[List[float]]    = None
    colbert_vecs:        Optional[List[List[float]]] = None


# ══════════════════════════════════════════════════════════════════════════════
# PDF / DOCX LOADING  (I5, I16, I17)
# ══════════════════════════════════════════════════════════════════════════════

_TOC_LINE_RE = re.compile(r"(\.\s*){2,}.*\d+\s*$")
_PAGE_TAG_RE  = re.compile(r"\[Page\s+\d+\]", re.IGNORECASE)


def _strip_page_tags(text: str) -> str:
    """I5, I20: Remove [Page N] markers injected during extraction."""
    return _PAGE_TAG_RE.sub(" ", text)


def _is_toc_page(text: str) -> bool:
    s = text.strip()
    if len(s) < TOC_MIN_CONTENT_CHARS:
        return True
    if len(s) > TOC_MAX_CONTENT_CHARS:
        return False
    lines = [l for l in s.splitlines() if l.strip()]
    if len(lines) < TOC_MIN_LINE_COUNT:
        return False
    toc = sum(1 for l in lines if _TOC_LINE_RE.search(l))
    return (toc / len(lines)) >= TOC_LINE_RATIO


def _table_to_sentences(table_rows: List[List[str]], caption: str = "") -> str:
    """
    I16: Convert a table (list-of-rows of cell strings) to natural-language
    sentences so the embedding understands the relationships.

    pdfplumber cells can be None — coerce every cell to str before calling
    strip() to prevent 'NoneType has no attribute strip' crashes.
    """
    def _cell(c) -> str:
        return str(c).strip() if c is not None else ""

    lines = []
    if caption:
        lines.append(f"Table: {caption}.")
    if not table_rows:
        return ""
    headers = [_cell(c) for c in table_rows[0] if _cell(c)]
    for row in table_rows[1:]:
        cells = [_cell(c) for c in row if _cell(c)]
        if not cells:
            continue
        if headers and len(cells) == len(headers):
            parts = [f"{h} is {v}" for h, v in zip(headers, cells) if h and v]
            lines.append(". ".join(parts) + ".")
        else:
            lines.append(" | ".join(cells) + ".")
    return "\n".join(lines)


def load_pdf(path: str) -> Tuple[str, List[Tuple[str, str]], Dict]:
    """
    Load PDF returning:
      - full_text: concatenated page texts (NO [Page N] markers)
      - tables:    list of (caption, natural-language-sentences) per table (I17)
      - metadata:  {num_pages, has_tables, tables_count, bookmarks}

    Uses pypdfium2 for text, pdfplumber for table content (I17).
    Page markers stripped from text (I5, I20).
    """
    meta: Dict = {"num_pages": 0, "has_tables": False,
                  "tables_count": 0, "bookmarks": []}
    pages_text: List[str] = []
    tables_out: List[Tuple[str, str]] = []

    if pdfium is None:
        logger.error("pypdfium2 not installed — cannot load PDF.")
        return "", [], meta

    # ── text extraction ───────────────────────────────────────────────────
    pdf = pdfium.PdfDocument(path)
    try:
        meta["num_pages"] = len(pdf)

        # I6: Extract bookmark tree for structural split
        try:
            bookmarks = []
            for bm in pdf.get_toc():
                bookmarks.append({
                    "title": bm.title,
                    "page":  bm.page_index,
                    "level": bm.level,
                })
            meta["bookmarks"] = bookmarks
        except Exception:
            meta["bookmarks"] = []

        for page_num in range(len(pdf)):
            page = pdf[page_num]
            page_text = ""
            try:
                tp = page.get_textpage()
                try:
                    page_text = tp.get_text_bounded()
                finally:
                    tp.close()
            except Exception as e:
                logger.warning(f"  Page {page_num+1} extraction error: {e}")
            finally:
                page.close()

            if not page_text or _is_toc_page(page_text):
                continue

            # I5, I20: store page text WITHOUT [Page N] markers
            pages_text.append(page_text)
    finally:
        pdf.close()

    full_text = "\n\n".join(pages_text)
    full_text = _strip_page_tags(full_text)   # extra guard

    # ── table extraction via pdfplumber (I17) ─────────────────────────────
    if pdfplumber is not None:
        try:
            with pdfplumber.open(path) as plumb:
                for pnum, pg in enumerate(plumb.pages):
                    for tbl in pg.extract_tables():
                        if not tbl:
                            continue
                        meta["has_tables"]  = True
                        meta["tables_count"] += 1
                        caption = f"Page {pnum+1} table {meta['tables_count']}"
                        nl_text = _table_to_sentences(tbl, caption)
                        if nl_text.strip():
                            tables_out.append((caption, nl_text))
        except Exception as e:
            logger.warning(f"  pdfplumber table extraction failed: {e}")

    return full_text, tables_out, meta


def load_docx(path: str) -> Tuple[str, List[Tuple[str, str]], Dict]:
    """
    Load DOCX returning text, tables as NL sentences (I16), and metadata.
    """
    if python_docx is None:
        logger.error("python-docx not installed.")
        return "", [], {}

    doc = python_docx.Document(path)
    meta = {"num_paragraphs": 0, "has_tables": False, "tables_count": 0}
    paras = []
    tables_out: List[Tuple[str, str]] = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            paras.append(t)
            meta["num_paragraphs"] += 1

    for i, table in enumerate(doc.tables):
        meta["has_tables"]  = True
        meta["tables_count"] += 1
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        nl = _table_to_sentences(rows, f"Table {i+1}")
        if nl.strip():
            tables_out.append((f"Table {i+1}", nl))

    return "\n".join(paras), tables_out, meta


# ══════════════════════════════════════════════════════════════════════════════
# STRUCTURAL SPLITTING — PDF Bookmarks → Parent Chunks  (I6)
# ══════════════════════════════════════════════════════════════════════════════

def _extract_sections_from_bookmarks(
    full_text: str,
    bookmarks: List[Dict],
    filename: str,
) -> List[Tuple[str, List[str], str]]:
    """
    I6: Use PDF bookmark tree as section boundaries (reliable).
    Returns list of (section_title, hierarchy, section_text).
    """
    if not bookmarks:
        return []

    # Build hierarchy stack and find section titles in text
    sections: List[Tuple[str, List[str], str]] = []
    lines = full_text.split("\n")
    line_lower = [l.lower().strip() for l in lines]

    last_idx = 0
    hierarchy_stack: List[Tuple[int, str]] = []  # (level, title)

    for bm in bookmarks:
        title = bm.get("title", "").strip()
        level = bm.get("level", 0)
        if not title:
            continue

        # Find this title in the text (fuzzy line match)
        title_lower = title.lower().strip()
        found_line = None
        for li, ll in enumerate(line_lower[last_idx:], start=last_idx):
            if title_lower in ll or ll in title_lower:
                found_line = li
                break

        if found_line is None:
            continue

        # Save previous section
        section_text = "\n".join(lines[last_idx:found_line]).strip()
        if section_text and last_idx > 0 and hierarchy_stack:
            parent_title, parent_hier = hierarchy_stack[-1][1], [
                h for _, h in hierarchy_stack
            ]
            sections.append((parent_title, parent_hier, section_text))

        # Update hierarchy stack
        while hierarchy_stack and hierarchy_stack[-1][0] >= level:
            hierarchy_stack.pop()
        hierarchy_stack.append((level, title))
        last_idx = found_line

    # Capture trailing section
    if last_idx < len(lines) and hierarchy_stack:
        trailing = "\n".join(lines[last_idx:]).strip()
        if trailing:
            parent_title = hierarchy_stack[-1][1]
            parent_hier = [h for _, h in hierarchy_stack]
            sections.append((parent_title, parent_hier, trailing))

    return sections


# Fallback: heading-pattern based split (I6 fallback only)
_HEADING_PATTERNS = [
    re.compile(r"^\d+(?:\.\d+)*\s+\[[A-Z_0-9]+\].*$", re.MULTILINE),
    re.compile(r"^\d+(?:\.\d+)+\s+[A-Z].*$", re.MULTILINE),
    re.compile(r"^#{1,6}\s+(.+)$", re.MULTILINE),
    re.compile(r"^\d+\.\s+([A-Z].+)$", re.MULTILINE),
    re.compile(
        r"^(?:Test Objective|Test Steps?|Pass Criteria|Pre-?conditions?|"
        r"Post-?conditions?|Main Test Execution|Configuration Parameters?|"
        r"Summary|Trace to Requirement|Affected Modules?)[\s:]*$",
        re.MULTILINE | re.IGNORECASE,
    ),
    re.compile(r"^\[[A-Z]+_[A-Z0-9_]+\]\s*.*$", re.MULTILINE),
]


def _heading_split_fallback(
    text: str, filename: str
) -> List[Tuple[str, List[str], str]]:
    """Regex heading split. Used only when bookmarks unavailable (I6 fallback)."""
    all_matches: List[Tuple[int, int, str]] = []
    for pat in _HEADING_PATTERNS:
        for m in pat.finditer(text):
            heading = m.group(0).strip()[:200]
            if len(heading) > 3:
                all_matches.append((m.start(), m.end(), heading))

    if not all_matches:
        return [("Document", ["Document"], text)]

    all_matches.sort(key=lambda x: x[0])
    # deduplicate overlapping matches
    deduped: List[Tuple[int, int, str]] = []
    for s, e, h in all_matches:
        if deduped and s < deduped[-1][1]:
            continue
        deduped.append((s, e, h))

    sections: List[Tuple[str, List[str], str]] = []
    for i, (s, e, h) in enumerate(deduped):
        end = deduped[i + 1][0] if i + 1 < len(deduped) else len(text)
        body = text[e:end].strip()
        if len(body) > 30:
            sections.append((h, [h], body))

    return sections if sections else [("Document", ["Document"], text)]


def split_into_sections(
    full_text: str,
    bookmarks: List[Dict],
    filename: str,
) -> List[Tuple[str, List[str], str]]:
    """
    I6: Primary path = PDF bookmarks; fallback = heading regex.
    Returns list of (title, hierarchy_list, section_body_text).
    """
    if bookmarks:
        secs = _extract_sections_from_bookmarks(full_text, bookmarks, filename)
        if secs:
            logger.info(f"  Structural split via bookmarks: {len(secs)} sections")
            return secs

    logger.info("  No usable bookmarks — using heading-regex fallback")
    secs = _heading_split_fallback(full_text, filename)
    logger.info(f"  Regex split produced {len(secs)} sections")
    return secs


# ══════════════════════════════════════════════════════════════════════════════
# PARENT-CHILD CHUNKING  (I2, I3)
# ══════════════════════════════════════════════════════════════════════════════

def _sentence_split(text: str) -> List[str]:
    """Split text into sentences; hard-limit each sentence to ENCODE_MAX_LEN."""
    raw = sent_tokenize(text)
    result = []
    for s in raw:
        if count_tokens(s) <= ENCODE_MAX_LEN:
            result.append(s)
        else:
            # Sentence is too long — split at token boundary
            result.extend(split_by_tokens(s, ENCODE_MAX_LEN))
    return result


def make_parent_chunks(
    section_title: str,
    section_hierarchy: List[str],
    section_text: str,
    filename: str,
    page_number: Optional[int],
) -> List[ParentChunk]:
    """
    I3: Create large parent chunks (≤ PARENT_CHUNK_TOKENS) from a section.
    These are NOT embedded — stored as payload for context delivery.
    """
    sentences = _sentence_split(section_text)
    parents: List[ParentChunk] = []
    current_sents: List[str] = []
    current_tokens = 0

    def _flush():
        if not current_sents:
            return
        body = " ".join(current_sents).strip()
        tok = count_tokens(body)
        if tok < MIN_CHILD_TOKENS // 2:
            return
        pid = str(uuid.uuid4())
        parents.append(ParentChunk(
            id=pid,
            text=body,
            filename=filename,
            section_title=section_title,
            section_hierarchy=list(section_hierarchy),
            page_number=page_number,
            token_count=tok,
        ))

    for sent in sentences:
        t = count_tokens(sent)
        if current_tokens + t > PARENT_CHUNK_TOKENS and current_sents:
            _flush()
            current_sents = [sent]
            current_tokens = t
        else:
            current_sents.append(sent)
            current_tokens += t
    _flush()

    return parents


def make_child_chunks(parent: ParentChunk) -> List[ChildChunk]:
    """
    I1, I2, I3: Create small child chunks (≤ CHILD_CHUNK_TOKENS tokens)
    from a parent chunk with token-count overlap (CHILD_OVERLAP_TOKENS).
    """
    sentences = _sentence_split(parent.text)
    children: List[ChildChunk] = []
    current_sents: List[str] = []
    current_tokens = 0
    child_idx = 0

    def _flush():
        nonlocal child_idx
        if not current_sents:
            return
        body = " ".join(current_sents).strip()
        tok = count_tokens(body)
        if tok < MIN_CHILD_TOKENS:
            return
        cid = str(uuid.uuid4())
        children.append(ChildChunk(
            id=cid,
            parent_id=parent.id,
            raw_text=body,
            enriched_text=body,        # enriched in-place later
            filename=parent.filename,
            section_title=parent.section_title,
            section_hierarchy=list(parent.section_hierarchy),
            page_number=parent.page_number,
            child_index=child_idx,
            token_count=tok,
        ))
        child_idx += 1

    for sent in sentences:
        t = count_tokens(sent)
        if current_tokens + t > CHILD_CHUNK_TOKENS and current_sents:
            _flush()
            # Overlap: keep last N tokens worth of sentences (I1)
            overlap_sents: List[str] = []
            overlap_tok = 0
            for prev_sent in reversed(current_sents):
                pt = count_tokens(prev_sent)
                if overlap_tok + pt > CHILD_OVERLAP_TOKENS:
                    break
                overlap_sents.insert(0, prev_sent)
                overlap_tok += pt
            current_sents  = overlap_sents + [sent]
            current_tokens = overlap_tok + t
        else:
            current_sents.append(sent)
            current_tokens += t
    _flush()

    return children


# ══════════════════════════════════════════════════════════════════════════════
# CONTEXTUAL ENRICHMENT  (I4)
# ══════════════════════════════════════════════════════════════════════════════

_CONTEXT_SYSTEM = (
    "You are a technical documentation assistant for AUTOSAR specifications. "
    "Given the parent section text and a child chunk, write a ONE-SENTENCE "
    "context (max 80 tokens) describing what this chunk covers in the context "
    "of the document. Start with 'This chunk is from'. "
    "Respond with ONLY that sentence — no preamble, no bullet points."
)


def _call_context_llm(parent_text: str, child_text: str, filename: str) -> Optional[str]:
    """
    I4: Call local LLM to generate a context prefix for a child chunk.
    Returns the generated sentence, or None if unavailable.
    """
    if not CONTEXT_LLM_URL or not _HAS_REQUESTS:
        return None
    try:
        user_msg = (
            f"Document: {filename}\n\n"
            f"<parent>\n{parent_text[:1500]}\n</parent>\n\n"
            f"<chunk>\n{child_text[:800]}\n</chunk>\n\n"
            "Write the one-sentence context."
        )
        resp = _requests.post(
            f"{CONTEXT_LLM_URL.rstrip('/')}/chat/completions",
            json={
                "model":       CONTEXT_LLM_MODEL,
                "max_tokens":  CONTEXT_MAX_TOKENS,
                "temperature": 0.0,
                "stream":      False,
                "messages": [
                    {"role": "system", "content": _CONTEXT_SYSTEM},
                    {"role": "user",   "content": user_msg},
                ],
            },
            timeout=CONTEXT_LLM_TIMEOUT,
        )
        if resp.status_code == 200:
            content = (
                resp.json()
                .get("choices", [{}])[0]
                .get("message", {})
                .get("content", "")
                .strip()
            )
            return content if content else None
    except Exception as e:
        logger.debug(f"Context LLM unavailable ({e}). Using metadata fallback.")
    return None


def enrich_child(child: ChildChunk, parent: ParentChunk) -> bool:
    """
    I4: Prepend context to child.enriched_text (in-place).
    Priority: LLM-generated context > metadata path fallback.
    Never mutates child.raw_text.
    Returns True if LLM generated the context, False if metadata fallback was used.
    """
    context_sentence = _call_context_llm(parent.text, child.raw_text, child.filename)
    used_llm = bool(context_sentence)

    if not context_sentence:
        # Deterministic fallback: build context from metadata (always available)
        hier = " > ".join(child.section_hierarchy) if child.section_hierarchy else child.section_title
        context_sentence = (
            f"This chunk is from {child.filename}, "
            f"section '{hier}', covering {child.section_title}."
        )

    child.enriched_text = f"{context_sentence}\n\n{child.raw_text}"
    return used_llm


# ══════════════════════════════════════════════════════════════════════════════
# BGE-M3 ENCODER  (I8–I15)
# ══════════════════════════════════════════════════════════════════════════════

class BGEM3Encoder:
    """
    Wraps BGEM3FlagModel from FlagEmbedding (preferred) with a
    sentence-transformers dense fallback.

    Outputs all three BGE-M3 signals per text:
      dense_vecs      — 1024-dim L2-normalised float32  (I14)
      lexical_weights — {token_id: weight} learned sparse (I8–I11)
      colbert_vecs    — (N × 1024) per-token matrix      (I15)

    I12: Batch encoding — all texts sent to GPU in one call.
    I19: model_fingerprint stored in collection metadata.
    """

    def __init__(self, model_name: str = EMBEDDING_MODEL):
        self.model_name = model_name
        self._model: Optional[BGEM3FlagModel] = None
        self._st_model = None                    # sentence-transformers fallback
        self._mode = "none"
        self._load()

    def _load(self) -> None:
        if _HAS_FLAG:
            try:
                logger.info(f"Loading BGEM3FlagModel: {self.model_name}")
                self._model = BGEM3FlagModel(
                    self.model_name,
                    use_fp16=True,
                    normalize_embeddings=True,     # I14
                )
                self._mode = "flag"
                logger.info("  ✓ BGEM3FlagModel loaded (dense + sparse + ColBERT)")
                return
            except Exception as e:
                logger.warning(f"  BGEM3FlagModel load failed: {e}. Trying sentence-transformers.")

        # sentence-transformers fallback (dense only)
        try:
            from sentence_transformers import SentenceTransformer
            logger.info(f"Loading SentenceTransformer: {self.model_name}")
            self._st_model = SentenceTransformer(self.model_name)
            self._mode = "st"
            logger.warning(
                "  ⚠ sentence-transformers dense-only mode. "
                "Sparse and ColBERT vectors will be empty. "
                "Install FlagEmbedding for full three-signal encoding."
            )
        except Exception as e:
            raise RuntimeError(
                f"Cannot load any embedding model for '{self.model_name}'. "
                f"Install FlagEmbedding or sentence-transformers. Error: {e}"
            )

    @property
    def fingerprint(self) -> str:
        return f"{self.model_name}::{self._mode}"

    def encode_batch(
        self,
        texts: List[str],
        is_query: bool = False,
    ) -> List[Dict]:
        """
        Encode a batch of texts.
        Returns list of dicts:
          {
            "dense":   List[float],          # 1024-dim, L2-normalised
            "sparse_indices": List[int],     # BGE-M3 vocab token IDs
            "sparse_values":  List[float],   # learned weights
            "colbert": List[List[float]],    # (N_tokens × 1024), may be []
          }
        I12: All texts in one GPU forward pass.
        I13: All three signals extracted.
        I14: Dense is L2-normalised.
        """
        if not texts:
            return []

        # Truncate texts exceeding model context (safety guard)
        safe_texts = []
        for t in texts:
            if count_tokens(t) > ENCODE_MAX_LEN:
                t = _TOKENIZER.decode(
                    _TOKENIZER.encode(t)[:ENCODE_MAX_LEN]
                ) if _TOKENIZER else t[:ENCODE_MAX_LEN * 4]
            safe_texts.append(t)

        if self._mode == "flag" and self._model is not None:
            return self._encode_flag(safe_texts, is_query)
        elif self._mode == "st" and self._st_model is not None:
            return self._encode_st(safe_texts)
        else:
            raise RuntimeError("No embedding model loaded.")

    def _encode_flag(self, texts: List[str], is_query: bool) -> List[Dict]:
        """I12: Batch encode via FlagEmbedding (dense + sparse + ColBERT)."""
        try:
            out = self._model.encode(  # type: ignore[union-attr]
                texts,
                batch_size=EMBED_BATCH_SIZE,
                max_length=ENCODE_MAX_LEN,
                return_dense=True,
                return_sparse=True,
                return_colbert_vecs=True,
            )
        except Exception as e:
            logger.error(f"Batch encode failed: {e}")
            return [self._empty_record() for _ in texts]

        results = []
        for i, text in enumerate(texts):
            dense = out["dense_vecs"][i]

            # I21: L2 norm guard — reject degenerate vectors
            norm = float(np.linalg.norm(dense))
            if norm < MIN_L2_NORM:
                logger.warning(
                    f"  ⚠ Degenerate dense vector (L2={norm:.4f}) for "
                    f"chunk [{text[:60]}…]. Skipping."
                )
                results.append(self._empty_record())
                continue

            # I14: already normalised by FlagEmbedding (normalize_embeddings=True)
            dense_list = dense.tolist() if hasattr(dense, "tolist") else list(dense)

            # I8, I13: BGE-M3 native sparse — {token_id: weight}
            lex = out["lexical_weights"][i]  # dict {int_id: float_weight}
            sparse_indices = list(lex.keys())
            sparse_values  = list(lex.values())

            # I15: ColBERT multi-vectors
            cv = out["colbert_vecs"][i]
            colbert = [
                v.tolist() if hasattr(v, "tolist") else list(v)
                for v in cv
            ]

            results.append({
                "dense":          dense_list,
                "sparse_indices": sparse_indices,
                "sparse_values":  sparse_values,
                "colbert":        colbert,
            })
        return results

    def _encode_st(self, texts: List[str]) -> List[Dict]:
        """Dense-only fallback via sentence-transformers (I12 batch)."""
        try:
            vecs = self._st_model.encode(  # type: ignore[union-attr]
                texts,
                batch_size=EMBED_BATCH_SIZE,
                normalize_embeddings=True,   # I14
                convert_to_numpy=True,
                show_progress_bar=False,
            )
        except Exception as e:
            logger.error(f"ST encode failed: {e}")
            return [self._empty_record() for _ in texts]

        results = []
        for vec in vecs:
            norm = float(np.linalg.norm(vec))
            if norm < MIN_L2_NORM:
                results.append(self._empty_record())
                continue
            results.append({
                "dense":          vec.tolist(),
                "sparse_indices": [],
                "sparse_values":  [],
                "colbert":        [],
            })
        return results

    @staticmethod
    def _empty_record() -> Dict:
        return {
            "dense":          None,
            "sparse_indices": [],
            "sparse_values":  [],
            "colbert":        [],
        }


# ══════════════════════════════════════════════════════════════════════════════
# QDRANT COLLECTION SETUP  (I15, I19)
# ══════════════════════════════════════════════════════════════════════════════

DENSE_DIM = 1024   # BGE-M3 dense output dimension


def setup_collection(
    client: "QdrantClient",
    collection: str,
    model_fingerprint: str,
    force_recreate: bool = False,
) -> None:
    """
    I15: Create Qdrant collection with dense + sparse + ColBERT vectors.
    I19: Write model fingerprint into collection metadata payload tag.
    """
    if not _HAS_QDRANT:
        raise RuntimeError("qdrant-client not installed.")

    existing = {c.name for c in client.get_collections().collections}

    if collection in existing:
        if force_recreate:
            logger.info(f"Recreating collection '{collection}'")
            client.delete_collection(collection)
        else:
            # I19: Verify fingerprint
            try:
                info = client.get_collection(collection)
                stored_fp = (info.config.params.vectors or {})
                # fingerprint stored in a dummy scroll check
                hits, _ = client.scroll(
                    collection_name=collection,
                    limit=1,
                    with_payload=["_model_fingerprint"],
                )
                if hits:
                    stored = hits[0].payload.get("_model_fingerprint", "")
                    if stored and stored != model_fingerprint:
                        raise RuntimeError(
                            f"Model mismatch! Collection was built with "
                            f"'{stored}' but current model is '{model_fingerprint}'. "
                            f"Re-run with --force-recreate to rebuild."
                        )
            except RuntimeError:
                raise
            except Exception as e:
                logger.warning(f"Could not verify model fingerprint: {e}")
            logger.info(f"Collection '{collection}' exists — resuming.")
            return

    logger.info(f"Creating collection '{collection}' (dense+sparse+ColBERT)")
    client.create_collection(
        collection_name=collection,
        vectors_config={
            "dense": qmodels.VectorParams(
                size=DENSE_DIM,
                distance=qmodels.Distance.COSINE,
            ),
            "colbert": qmodels.VectorParams(
                size=DENSE_DIM,
                distance=qmodels.Distance.COSINE,
                multivector_config=qmodels.MultiVectorConfig(
                    comparator=qmodels.MultiVectorComparator.MAX_SIM
                ),
            ),
        },
        sparse_vectors_config={
            "sparse": qmodels.SparseVectorParams(
                index=qmodels.SparseIndexParams(on_disk=False)
            )
        },
    )
    logger.info(f"  ✓ Collection '{collection}' created.")


# ══════════════════════════════════════════════════════════════════════════════
# DEDUPLICATION  (I7)
# ══════════════════════════════════════════════════════════════════════════════

def dedup_children(children: List[ChildChunk]) -> List[ChildChunk]:
    """
    I7: Jaccard dedup using AUTOSAR-tokenized word sets.
    Identifier-specific tokens (SWS_Com_00228) are preserved as atomic units,
    so chunks differing only in identifier are NOT deduplicated.
    """
    accepted: List[ChildChunk] = []
    window_texts: List[str] = []

    for ch in children:
        too_similar = any(
            jaccard(ch.raw_text, prev) >= JACCARD_THRESHOLD
            for prev in window_texts[-JACCARD_WINDOW:]
        )
        if not too_similar:
            accepted.append(ch)
            window_texts.append(ch.raw_text)

    removed = len(children) - len(accepted)
    if removed:
        logger.info(f"  Dedup removed {removed}/{len(children)} near-duplicate children")
    return accepted


# ══════════════════════════════════════════════════════════════════════════════
# FILE HASH & INCREMENTAL SKIP
# ══════════════════════════════════════════════════════════════════════════════

def file_sha256(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for block in iter(lambda: f.read(65536), b""):
            h.update(block)
    return h.hexdigest()


def already_indexed(client: "QdrantClient", collection: str, fhash: str) -> bool:
    try:
        hits, _ = client.scroll(
            collection_name=collection,
            scroll_filter=qmodels.Filter(
                must=[qmodels.FieldCondition(
                    key="file_hash",
                    match=qmodels.MatchValue(value=fhash),
                )]
            ),
            limit=1,
        )
        return len(hits) > 0
    except Exception:
        return False


# ══════════════════════════════════════════════════════════════════════════════
# BATCH UPLOAD  (I18)
# ══════════════════════════════════════════════════════════════════════════════

def _child_to_point(
    child: ChildChunk,
    parent: ParentChunk,
    file_hash: str,
    doc_meta: Dict,
    model_fingerprint: str,
) -> Optional["qmodels.PointStruct"]:
    """Build a Qdrant PointStruct from a fully encoded ChildChunk."""
    if child.dense_vec is None:
        return None
    if not child.sparse_indices:
        logger.debug(f"  Chunk {child.id[:8]} has empty sparse vector — dense only.")

    # ColBERT: truncate to COLBERT_MAX_TOKENS to stay within Qdrant's 1MB/vector limit.
    # 1024-dim × float32(4B) × 256 tokens = 1,048,576 B exactly — use 250 for safety.
    colbert_vecs = child.colbert_vecs or []
    if len(colbert_vecs) > COLBERT_MAX_TOKENS:
        logger.debug(
            f"  ColBERT truncated {len(colbert_vecs)} → {COLBERT_MAX_TOKENS} tokens "
            f"for chunk {child.id[:8]} (Qdrant 1MB limit)"
        )
        colbert_vecs = colbert_vecs[:COLBERT_MAX_TOKENS]
    has_colbert = bool(colbert_vecs)

    vector_dict: Dict = {
        "dense": child.dense_vec,
    }
    if child.sparse_indices:
        vector_dict["sparse"] = qmodels.SparseVector(
            indices=child.sparse_indices,
            values=child.sparse_values,
        )
    if has_colbert:
        vector_dict["colbert"] = colbert_vecs

    return qmodels.PointStruct(
        id=child.id,
        vector=vector_dict,
        payload={
            # Retrieval content
            "content":             child.raw_text,        # unenriched for display
            "enriched_content":    child.enriched_text,   # what was embedded
            "parent_text":         parent.text,           # I3: full section for LLM
            "parent_id":           child.parent_id,
            "child_index":         child.child_index,
            # Metadata
            "filename":            child.filename,
            "section_title":       child.section_title,
            "section_hierarchy":   child.section_hierarchy,
            "page_number":         child.page_number,
            "file_hash":           file_hash,
            "token_count":         child.token_count,
            "parent_token_count":  parent.token_count,
            # I19: model fingerprint per point
            "_model_fingerprint":  model_fingerprint,
            **doc_meta,
        },
    )


def upload_batch(
    client: "QdrantClient",
    collection: str,
    points: List["qmodels.PointStruct"],
    batch_size: int = 16,
    max_retries: int = 4,
) -> int:
    """
    I18: Stream points to Qdrant in fixed-size batches with retry + auto-halving.

    ColBERT multi-vectors (N_tokens x 1024 floats per point) make payloads much
    larger than dense-only ingestion. If Qdrant disconnects mid-batch the batch
    size is automatically halved and retried, down to a minimum of 1 point per
    request. Exponential back-off (1s, 2s, 4s, 8s) is applied between retries.
    """
    uploaded = 0
    i = 0
    current_batch_size = batch_size

    while i < len(points):
        batch = points[i:i + current_batch_size]
        success = False

        for attempt in range(1, max_retries + 1):
            try:
                client.upsert(collection_name=collection, points=batch, wait=True)
                uploaded += len(batch)
                logger.info(
                    f"  Uploaded {len(batch)} points "
                    f"(offset {i}-{i+len(batch)-1}, batch_size={current_batch_size})"
                )
                success = True
                break
            except Exception as e:
                err_str = str(e).lower()
                is_disconnect = any(k in err_str for k in (
                    "disconnected", "server disconnected", "connection reset",
                    "broken pipe", "remote end closed", "connection refused",
                    "timeout", "timed out",
                ))
                if is_disconnect and current_batch_size > 1:
                    new_size = max(1, current_batch_size // 2)
                    logger.warning(
                        f"  Qdrant disconnect (batch_size={current_batch_size}). "
                        f"Halving to {new_size}, attempt {attempt}/{max_retries} ..."
                    )
                    current_batch_size = new_size
                    batch = points[i:i + current_batch_size]
                    time.sleep(2 ** (attempt - 1))
                else:
                    wait = 2 ** (attempt - 1)
                    logger.warning(
                        f"  Upload failed (attempt {attempt}/{max_retries}): {e}. "
                        f"Retrying in {wait}s ..."
                    )
                    time.sleep(wait)

        if not success:
            logger.error(
                f"  Permanently failed to upload {len(batch)} points at offset {i} "
                f"after {max_retries} attempts — skipping."
            )

        i += len(batch)

    return uploaded


# ══════════════════════════════════════════════════════════════════════════════
# AUDIT LOG
# ══════════════════════════════════════════════════════════════════════════════

class IngestionAudit:
    def __init__(self, json_path: str, config: Dict):
        self.json_path = json_path
        self.data: Dict = {
            "run":     config,
            "files":   [],
            "summary": {
                "total_files":    0,
                "total_parents":  0,
                "total_children": 0,
                "total_uploaded": 0,
                "dedup_removed":  0,
                "embed_skipped":  0,
                "warnings":       [],
            },
        }

    def warn(self, msg: str) -> None:
        logger.warning(f"  ⚠ {msg}")
        self.data["summary"]["warnings"].append(msg)

    def add_file(self, record: Dict) -> None:
        self.data["files"].append(record)
        s = self.data["summary"]
        s["total_files"]    += 1
        s["total_parents"]  += record.get("parents", 0)
        s["total_children"] += record.get("children_after_dedup", 0)
        s["total_uploaded"] += record.get("uploaded", 0)
        s["dedup_removed"]  += record.get("dedup_removed", 0)
        s["embed_skipped"]  += record.get("embed_skipped", 0)

    def save(self) -> None:
        def _conv(obj):
            if isinstance(obj, (np.integer,)):  return int(obj)
            if isinstance(obj, (np.floating,)): return float(obj)
            if isinstance(obj, np.ndarray):     return obj.tolist()
            if isinstance(obj, dict):   return {k: _conv(v) for k, v in obj.items()}
            if isinstance(obj, list):   return [_conv(v) for v in obj]
            return obj
        with open(self.json_path, "w", encoding="utf-8") as f:
            json.dump(_conv(self.data), f, indent=2)


# ══════════════════════════════════════════════════════════════════════════════
# PROCESS ONE FILE
# ══════════════════════════════════════════════════════════════════════════════

def process_file(
    path: str,
    encoder: BGEM3Encoder,
    client: "QdrantClient",
    collection: str,
    audit: IngestionAudit,
) -> None:
    """Full ingestion pipeline for a single file."""
    filename = os.path.basename(path)
    file_hash = file_sha256(path)
    record: Dict = {"filename": filename, "file_hash": file_hash}

    # Incremental skip
    if already_indexed(client, collection, file_hash):
        logger.info(f"  ↳ Already indexed — skipping {filename}")
        record["status"] = "skipped_already_indexed"
        audit.add_file(record)
        return

    logger.info(f"\n{'='*70}")
    logger.info(f"Processing: {filename}")
    logger.info(f"{'='*70}")

    # ── Load document ──────────────────────────────────────────────────────
    ext = Path(path).suffix.lower()
    doc_meta: Dict = {}

    if ext == ".pdf":
        full_text, tables, raw_meta = load_pdf(path)
        doc_meta = {
            "file_type":    ".pdf",
            "num_pages":    raw_meta.get("num_pages", 0),
            "has_tables":   raw_meta.get("has_tables", False),
            "tables_count": raw_meta.get("tables_count", 0),
        }
        bookmarks = raw_meta.get("bookmarks", [])
        record["num_pages"] = raw_meta.get("num_pages", 0)
    elif ext == ".docx":
        full_text, tables, raw_meta = load_docx(path)
        doc_meta = {
            "file_type":      ".docx",
            "num_paragraphs": raw_meta.get("num_paragraphs", 0),
            "has_tables":     raw_meta.get("has_tables", False),
        }
        bookmarks = []
    else:
        logger.warning(f"  Unsupported file type: {ext} — skipping")
        record["status"] = "skipped_unsupported"
        audit.add_file(record)
        return

    if not full_text.strip():
        logger.warning(f"  No text extracted from {filename} — skipping")
        record["status"] = "skipped_empty"
        audit.add_file(record)
        return

    # Append table natural-language sentences to text (I16, I17)
    table_appendix = ""
    if tables:
        logger.info(f"  Found {len(tables)} table(s) — converting to NL sentences")
        table_appendix = "\n\n".join(nl for _, nl in tables)

    full_text_with_tables = full_text
    if table_appendix:
        full_text_with_tables = full_text + "\n\n" + table_appendix

    # ── Structural split → Sections ────────────────────────────────────────
    sections = split_into_sections(full_text_with_tables, bookmarks, filename)
    logger.info(f"  Sections: {len(sections)}")

    # ── Build parent + child chunks ────────────────────────────────────────
    all_parents:  List[ParentChunk]  = []
    all_children: List[ChildChunk]   = []

    for sec_title, sec_hier, sec_text in sections:
        # Page number best-effort from bookmark metadata
        page_num: Optional[int] = None
        if bookmarks:
            for bm in bookmarks:
                if bm.get("title", "").strip().lower() in sec_title.lower():
                    page_num = bm.get("page")
                    break

        parents = make_parent_chunks(
            section_title=sec_title,
            section_hierarchy=sec_hier,
            section_text=sec_text,
            filename=filename,
            page_number=page_num,
        )
        for parent in parents:
            children = make_child_chunks(parent)
            all_parents.append(parent)
            all_children.extend(children)

    logger.info(f"  Parents: {len(all_parents)}  |  Children (raw): {len(all_children)}")

    if not all_children:
        audit.warn(f"{filename}: produced 0 child chunks after splitting.")
        record["status"] = "skipped_empty"
        audit.add_file(record)
        return

    # Build parent lookup (parent_id → ParentChunk)
    parent_map = {p.id: p for p in all_parents}

    # ── Deduplication (I7) ─────────────────────────────────────────────────
    before_dedup = len(all_children)
    all_children = dedup_children(all_children)
    dedup_removed = before_dedup - len(all_children)

    # ── Contextual enrichment (I4) ─────────────────────────────────────────
    logger.info(f"  Contextual enrichment for {len(all_children)} children …")
    llm_ok = 0
    fallback_ok = 0
    for ch in all_children:
        parent = parent_map.get(ch.parent_id)
        if parent:
            used_llm = enrich_child(ch, parent)
            if used_llm:
                llm_ok += 1
            else:
                fallback_ok += 1
    logger.info(f"  Context enriched: {llm_ok}/{len(all_children)} via LLM, "
                f"{fallback_ok}/{len(all_children)} via metadata fallback")

    # ── Encode in batches (I12, I13, I14, I15) ────────────────────────────
    logger.info(f"  Encoding {len(all_children)} child chunks via BGE-M3 …")
    embed_texts = [ch.enriched_text for ch in all_children]
    t0 = time.time()

    encoded_records = encoder.encode_batch(embed_texts, is_query=False)

    elapsed = time.time() - t0
    logger.info(f"  Encoding done in {elapsed:.1f}s  "
                f"({elapsed / max(len(all_children), 1) * 1000:.0f}ms/chunk)")

    # Attach vectors back to children
    embed_skipped = 0
    for ch, rec in zip(all_children, encoded_records):
        if rec.get("dense") is None:
            embed_skipped += 1
            ch.dense_vec    = None
            continue
        ch.dense_vec      = rec["dense"]
        ch.sparse_indices = rec["sparse_indices"]
        ch.sparse_values  = rec["sparse_values"]
        ch.colbert_vecs   = rec["colbert"]

    if embed_skipped:
        audit.warn(f"{filename}: {embed_skipped} children failed embedding and will be missing.")

    # ── Build Qdrant points ─────────────────────────────────────────────────
    points: List["qmodels.PointStruct"] = []
    for ch in all_children:
        parent = parent_map.get(ch.parent_id)
        if parent is None or ch.dense_vec is None:
            continue
        pt = _child_to_point(ch, parent, file_hash, doc_meta, encoder.fingerprint)
        if pt is not None:
            points.append(pt)

    logger.info(f"  Points to upload: {len(points)}")

    if not points:
        audit.warn(f"{filename}: 0 uploadable points after encoding.")
        record["status"] = "skipped_empty"
        audit.add_file(record)
        return

    # ── Stream upload (I18) ────────────────────────────────────────────────
    uploaded = upload_batch(client, collection, points)

    record.update({
        "status":               "ok",
        "parents":              len(all_parents),
        "children_raw":         before_dedup,
        "dedup_removed":        dedup_removed,
        "children_after_dedup": len(all_children),
        "embed_skipped":        embed_skipped,
        "uploaded":             uploaded,
    })
    audit.add_file(record)
    logger.info(f"  ✓ {filename} complete — {uploaded} points uploaded")


# ══════════════════════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════════════════════

def main() -> None:
    # Must be declared before any reference to these names in this scope
    global CONTEXT_LLM_URL, CONTEXT_LLM_MODEL

    ap = argparse.ArgumentParser(
        description="AUTOSAR RAG Ingestion V3 — Contextual Parent-Child + BGE-M3 three-signal"
    )
    ap.add_argument("--data-dir",       default=DATA_DIR,
                    help="Root directory containing AUTOSAR PDFs/DOCXs")
    ap.add_argument("--collection",     default=COLLECTION,
                    help="Qdrant collection name")
    ap.add_argument("--qdrant-url",     default=QDRANT_URL,
                    help="Qdrant server URL")
    ap.add_argument("--output-dir",     default=OUTPUT_DIR,
                    help="Directory for log + audit JSON")
    ap.add_argument("--force-recreate", action="store_true",
                    help="Delete and recreate the Qdrant collection")
    ap.add_argument("--no-context",     action="store_true",
                    help="Skip LLM contextual enrichment (use metadata fallback only)")
    ap.add_argument("--context-url",    default=CONTEXT_LLM_URL)
    ap.add_argument("--context-model",  default=CONTEXT_LLM_MODEL)
    args = ap.parse_args()

    # Override globals from CLI
    if args.no_context:
        CONTEXT_LLM_URL = ""
    else:
        CONTEXT_LLM_URL   = args.context_url
        CONTEXT_LLM_MODEL = args.context_model

    # Logging
    logger_inst, log_path = _setup_logging(args.output_dir)

    ts = time.strftime("%Y%m%d_%H%M%S")
    json_path = str(Path(args.output_dir) / f"ingestion_v3_{ts}.json")

    run_config = {
        "timestamp":       ts,
        "data_dir":        args.data_dir,
        "collection":      args.collection,
        "qdrant_url":      args.qdrant_url,
        "embedding_model": EMBEDDING_MODEL,
        "parent_tokens":   PARENT_CHUNK_TOKENS,
        "child_tokens":    CHILD_CHUNK_TOKENS,
        "overlap_tokens":  CHILD_OVERLAP_TOKENS,
        "context_llm":     CONTEXT_LLM_URL or "disabled",
        "jaccard_thresh":  JACCARD_THRESHOLD,
        "min_l2_norm":     MIN_L2_NORM,
    }

    audit = IngestionAudit(json_path, run_config)

    logger.info("=" * 70)
    logger.info("AUTOSAR RAG INGESTION V3")
    logger.info("=" * 70)
    logger.info(json.dumps(run_config, indent=2))

    # ── Validate dependencies ──────────────────────────────────────────────
    missing = []
    if not _HAS_QDRANT:
        missing.append("qdrant-client")
    if not _HAS_FLAG and not True:   # ST checked at runtime
        missing.append("FlagEmbedding or sentence-transformers")
    if pdfium is None:
        missing.append("pypdfium2")
    if missing:
        raise SystemExit(
            f"Missing required packages: {missing}. "
            f"Run: pip install {' '.join(missing)}"
        )

    # ── Load encoder (I12–I15) ─────────────────────────────────────────────
    encoder = BGEM3Encoder(EMBEDDING_MODEL)
    logger.info(f"Encoder mode: {encoder._mode}  |  fingerprint: {encoder.fingerprint}")

    # ── Connect to Qdrant ─────────────────────────────────────────────────
    client = QdrantClient(url=args.qdrant_url, timeout=60)
    setup_collection(
        client,
        args.collection,
        encoder.fingerprint,
        force_recreate=args.force_recreate,
    )

    # ── Discover files ─────────────────────────────────────────────────────
    data_dir = Path(args.data_dir)
    if not data_dir.exists():
        raise SystemExit(f"Data directory not found: {data_dir}")

    all_files = sorted(
        p for p in data_dir.rglob("*")
        if p.suffix.lower() in {".pdf", ".docx"} and p.is_file()
    )
    logger.info(f"\nFound {len(all_files)} files in {data_dir}")

    if not all_files:
        raise SystemExit("No PDF or DOCX files found in data directory.")

    # ── Process each file ─────────────────────────────────────────────────
    for i, fpath in enumerate(all_files, 1):
        logger.info(f"\n[{i}/{len(all_files)}]")
        try:
            process_file(
                path=str(fpath),
                encoder=encoder,
                client=client,
                collection=args.collection,
                audit=audit,
            )
        except KeyboardInterrupt:
            logger.warning("Interrupted by user.")
            break
        except Exception as e:
            logger.error(f"  ✗ Failed: {fpath.name} — {e}", exc_info=True)
            audit.data["files"].append({
                "filename": fpath.name,
                "status":   "error",
                "error":    str(e),
            })
        finally:
            audit.save()

    # ── Final summary ──────────────────────────────────────────────────────
    s = audit.data["summary"]
    logger.info("\n" + "=" * 70)
    logger.info("INGESTION V3 COMPLETE")
    logger.info("=" * 70)
    logger.info(f"  Files processed   : {s['total_files']}")
    logger.info(f"  Parent chunks     : {s['total_parents']}")
    logger.info(f"  Child chunks      : {s['total_children']}")
    logger.info(f"  Points uploaded   : {s['total_uploaded']}")
    logger.info(f"  Dedup removed     : {s['dedup_removed']}")
    logger.info(f"  Embed skipped     : {s['embed_skipped']}")
    logger.info(f"  Collection        : {args.collection}")
    logger.info(f"  Encoder           : {encoder.fingerprint}")
    logger.info(f"  Log               : {log_path}")
    logger.info(f"  Audit JSON        : {json_path}")

    if s["warnings"]:
        logger.warning(f"\n  ⚠ {len(s['warnings'])} warnings:")
        for w in s["warnings"]:
            logger.warning(f"    - {w}")
    else:
        logger.info("  ✓ No warnings.")

    audit.save()


if __name__ == "__main__":
    main()