# """
# ADVANCED RAG INGESTION SYSTEM V2 (CLEAN)
# =========================================
# Improvements over V1:
# - pypdfium2 for PDF text extraction (proper word spacing, no concatenated table words)
# - TOC / boilerplate page detection and skipping
#   * Ratio threshold raised 0.30 → 0.50 (need majority of lines to match)
#   * Minimum 5 non-empty lines required before ratio check (prevents false positives
#     on short section-header pages like "Version ......... 3.0.0")
#   * Max-chars guard: pages > 800 chars are never discarded regardless of ratio
#   * Skipped page numbers are logged for auditability
# - Near-duplicate chunk removal via Jaccard similarity (85% threshold)
# - Larger chunk overlap (256 chars instead of 128) to reduce boundary failures
# - Embedding robustness: pre-truncate chunks > MAX_EMBED_CHARS before sending to
#   Ollama; on HTTP 500, truncate and retry instead of silently dropping the chunk;
#   per-file summary log of any remaining skips
# - All V1 features preserved: section-aware chunking, BM25, hybrid Qdrant upload,
#   OllamaBGEM3Embedder with retry, deduplication by file hash, dimension-mismatch check
# """

# import argparse
# import json
# import math
# import numpy as np
# import os
# import re
# import hashlib
# import uuid
# import time
# from typing import Optional, List, Dict, Tuple
# from dataclasses import dataclass
# from pathlib import Path
# import logging

# import pypdfium2 as pdfium
# import pdfplumber          # kept for PDF table-count metadata only (not text extraction)
# import docx
# import requests
# from sentence_transformers import SentenceTransformer
# from qdrant_client import QdrantClient
# from qdrant_client.models import (
#     VectorParams,
#     Distance,
#     Filter,
#     FieldCondition,
#     MatchValue,
#     PointStruct,
#     SparseVector,
#     SparseVectorParams,
#     SparseIndexParams,
# )
# from rank_bm25 import BM25Okapi
# import nltk
# from nltk.tokenize import sent_tokenize, word_tokenize

# # Download NLTK data if needed
# try:
#     nltk.data.find('tokenizers/punkt')
# except LookupError:
#     try:
#         nltk.download('punkt_tab', quiet=True)
#     except Exception:
#         nltk.download('punkt', quiet=True)

# # ================= CONFIG =================

# DATA_DIR = r"/home/olj3kor/praveen/Autosar_docs_2"
# COLLECTION = "Autosar_dummy"
# QDRANT_URL = "http://localhost:7333"

# # Embedding options
# USE_OLLAMA_BGE_M3 = True
# OLLAMA_URL = "http://localhost:11434"
# # OLLAMA_URL = "http://localhost:8011/v1"
# # OLLAMA_MODEL = "bge-large:latest"
# OLLAMA_MODEL = "bge-m3:latest"

# # Fallback to SentenceTransformer if Ollama unavailable
# FALLBACK_MODEL = "sentence-transformers/all-MiniLM-L6-v2"

# # Chunking parameters
# CHUNK_SIZE = 1024
# CHUNK_OVERLAP = 256          # increased from 128 to reduce boundary-split failures
# MIN_CHUNK_SIZE = 100
# MAX_CHUNK_SIZE = 1024

# # Section detection
# ENABLE_SECTION_AWARE = True
# SECTION_PATTERNS = [
#     r'^#{1,6}\s+(.+)$',       # Markdown headers
#     r'^([A-Z][^.!?]*):$',     # Title case with colon
#     r'^\d+\.\s+([A-Z].+)$',   # Numbered sections
#     r'^([A-Z\s]{3,})$',       # All-caps headers (min 3 chars)
# ]

# BM25_OUTPUT = "bm25_index.json"

# # Near-duplicate suppression
# JACCARD_DEDUP_THRESHOLD = 0.90   # chunks with >85% Jaccard similarity are dropped
# JACCARD_WINDOW = 20              # compare each new chunk against last N accepted chunks

# # TOC detection
# TOC_LINE_RATIO = 0.50            # >50% of non-empty lines look like TOC entries → skip page
# TOC_MIN_CONTENT_CHARS = 50       # pages with fewer characters are classified as TOC/boilerplate
# TOC_MIN_LINE_COUNT = 5           # don't ratio-classify pages with fewer non-empty lines (section headers)
# TOC_MAX_CONTENT_CHARS = 800      # never classify dense pages as TOC regardless of ratio

# # Embedding safety
# MAX_EMBED_CHARS = 512           # truncate chunk text to this before sending to Ollama (fits any context window)

# # ----- RAGAS HeadlineSplitter chunker config (used only when --chunker ragas) -----
# # These mirror the settings in generate_dataset_v2.py so chunks produced here
# # match the reference_contexts in the dataset.
# NVIDIA_LLM_MODEL       = "meta/llama-3.1-70b-instruct"
# NVIDIA_BASE_URL        = "https://integrate.api.nvidia.com/v1"

# # If you used a local vLLM to generate the dataset, match those values here.
# VLLM_LLM_MODEL         = "Qwen/Qwen2.5-32B-Instruct-AWQ"
# VLLM_BASE_URL          = "http://localhost:8011/v1"

# # Embedding model RAGAS uses internally (only needed for the KG transform
# # pipeline — does NOT affect what Ollama does downstream).
# RAGAS_INTERNAL_EMBED_MODEL = "BAAI/bge-m3"

# # =========================================

# logging.basicConfig(
#     level=logging.INFO,
#     format='%(asctime)s - %(levelname)s - %(message)s'
# )
# logger = logging.getLogger(__name__)


# # ---------------------------------------------------------------------------
# # Dataclasses
# # ---------------------------------------------------------------------------

# @dataclass
# class DocumentSection:
#     """Represents a document section with hierarchy."""
#     title: str
#     content: str
#     level: int
#     page_number: Optional[int] = None
#     section_type: str = "text"
#     section_hierarchy: Optional[List[str]] = None

#     def __post_init__(self):
#         if self.section_hierarchy is None:
#             self.section_hierarchy = [self.title]


# @dataclass
# class EnrichedChunk:
#     """Chunk with rich metadata."""
#     text: str
#     section_title: str
#     section_hierarchy: List[str]
#     page_number: Optional[int]
#     chunk_type: str
#     word_count: int
#     sentence_count: int
#     start_char: int
#     end_char: int


# # ---------------------------------------------------------------------------
# # Helper functions
# # ---------------------------------------------------------------------------

# def jaccard_similarity(text_a: str, text_b: str) -> float:
#     """Compute word-level Jaccard similarity between two text strings."""
#     words_a = set(text_a.lower().split())
#     words_b = set(text_b.lower().split())
#     if not words_a or not words_b:
#         return 0.0
#     return len(words_a & words_b) / len(words_a | words_b)


# def file_hash(path: str) -> str:
#     """Generate SHA-256 hash of a file using block-wise reading."""
#     h = hashlib.sha256()
#     with open(path, "rb") as f:
#         for block in iter(lambda: f.read(65536), b""):
#             h.update(block)
#     return h.hexdigest()


# def already_indexed(client: QdrantClient, collection: str, file_hash_value: str) -> bool:
#     """Return True if the file hash is already present in the collection."""
#     try:
#         filt = Filter(
#             must=[FieldCondition(key="file_hash", match=MatchValue(value=file_hash_value))]
#         )
#         points, _ = client.scroll(collection_name=collection, scroll_filter=filt, limit=1)
#         return len(points) > 0
#     except Exception as e:
#         logger.warning(f"Could not check if file is already indexed: {e}")
#         return False


# # ---------------------------------------------------------------------------
# # Embedding models
# # ---------------------------------------------------------------------------

# class OllamaBGEM3Embedder:
#     """BGE-M3 embedder using Ollama with retry logic."""

#     def __init__(self, base_url: str = "http://localhost:11434", model: str = "bge-m3"):
#         self.base_url = base_url
#         self.model = model
#         self.dimension = 1024
#         self._test_connection()

#     def _test_connection(self):
#         """Test Ollama connection and verify model availability."""
#         try:
#             response = requests.get(f"{self.base_url}/api/tags", timeout=5)
#             if response.status_code == 200:
#                 logger.info(f"✓ Connected to Ollama at {self.base_url}")
#                 available_models = [m.get("name", "") for m in response.json().get("models", [])]
#                 if self.model not in available_models:
#                     logger.warning(
#                         f"Model '{self.model}' not found in Ollama. "
#                         f"Available: {available_models}"
#                     )
#             else:
#                 raise ConnectionError("Ollama not responding")
#         except Exception as e:
#             logger.error(f"✗ Cannot connect to Ollama: {e}")
#             raise

#     def encode(self, texts: List[str], batch_size: int = 8,
#                show_progress_bar: bool = False) -> List[Optional[List[float]]]:
#         """
#         Encode texts using Ollama BGE-M3 with per-text retry and automatic
#         pre-truncation for oversized inputs.

#         BGE-M3 in Ollama returns HTTP 500 when the prompt exceeds the model's
#         context window.  Every text is pre-truncated to MAX_EMBED_CHARS before
#         the first request — this eliminates context-window 500s entirely without
#         consuming a retry slot.  The three retry attempts are reserved for
#         genuine transient server or network errors.

#         On any failure the Ollama response body is included in the warning log
#         so the root cause is visible without inspecting Ollama's own log file.
#         A per-batch summary warns if any chunks end up missing from the index.
#         """
#         embeddings: List[Optional[List[float]]] = []
#         skipped_count = 0
#         patched_count = 0  # chunks where NaN/Inf components were replaced with 0.0

#         for i in range(0, len(texts), batch_size):
#             batch = texts[i:i + batch_size]

#             for text in batch:
#                 # Pre-truncate BEFORE the first request.  MAX_EMBED_CHARS (~1 000 tokens)
#                 # is well within BGE-M3's 8192-token context window, so this never
#                 # loses meaningful content and prevents all context-overflow 500s.
#                 safe_text = text if len(text) <= MAX_EMBED_CHARS else text[:MAX_EMBED_CHARS]

#                 embedding = None
#                 for attempt in range(3):
#                     try:
#                         response = requests.post(
#                             f"{self.base_url}/api/embeddings",
#                             json={"model": self.model, "prompt": safe_text},
#                             timeout=60
#                         )
#                         if response.status_code == 200:
#                             data = response.json()
#                             if "embedding" in data:
#                                 raw_vec = data["embedding"]
#                                 # Fast-path: math.isfinite check adds no measurable
#                                 # overhead and avoids numpy allocation on the common
#                                 # case where every component is a normal float.
#                                 if not all(math.isfinite(x) for x in raw_vec):
#                                     # Rare path: model produced NaN/Inf.  Replace
#                                     # non-finite components with 0.0 so the vector
#                                     # remains usable in Qdrant, and track how often
#                                     # this happens so operators can spot patterns.
#                                     arr = np.array(raw_vec, dtype=np.float64)
#                                     non_finite_mask = ~np.isfinite(arr)
#                                     n_bad = int(non_finite_mask.sum())
#                                     logger.warning(
#                                         f"Embedding attempt {attempt + 1}: "
#                                         f"{n_bad} non-finite component(s) (NaN/Inf) "
#                                         f"replaced with 0.0 in vector for chunk of "
#                                         f"{len(safe_text)} chars. "
#                                         f"Consider investigating this input text."
#                                     )
#                                     arr[non_finite_mask] = 0.0
#                                     raw_vec = arr.tolist()
#                                     patched_count += 1
#                                 embedding = raw_vec
#                                 break
#                             else:
#                                 # Ollama API changed key name or returned an error body
#                                 logger.warning(
#                                     f"Embedding attempt {attempt + 1}: "
#                                     f"unexpected response format (no 'embedding' key): "
#                                     f"{response.text[:120]}"
#                                 )
#                         else:
#                             body = response.text
#                             logger.warning(
#                                 f"Embedding attempt {attempt + 1} failed with status "
#                                 f"{response.status_code}: {body[:120]}"
#                             )
#                             # "json: unsupported value" means the model produced a
#                             # NaN/Inf that Go's JSON encoder refuses to serialise.
#                             # This is deterministic — the same text always triggers
#                             # the same failure, so retrying wastes time.  Break
#                             # immediately and let the chunk be skipped.
#                             if "unsupported value" in body:
#                                 logger.warning(
#                                     f"Embedding attempt {attempt + 1}: Ollama returned "
#                                     f"'json: unsupported value' (NaN/Inf in model output) "
#                                     f"— skipping retries for this chunk."
#                                 )
#                                 break
#                     except Exception as e:
#                         logger.warning(f"Embedding attempt {attempt + 1} error: {e}")
#                     if attempt < 2:
#                         time.sleep(2 ** attempt)

#                 if embedding is None:
#                     skipped_count += 1
#                     logger.warning(
#                         f"Failed to embed chunk after 3 attempts "
#                         f"(sent {len(safe_text)} chars, original {len(text)} chars). "
#                         f"Total skipped so far: {skipped_count}"
#                     )
#                 embeddings.append(embedding)

#             if show_progress_bar and (i // batch_size) % 10 == 0:
#                 logger.info(f"Encoded {min(i + batch_size, len(texts))}/{len(texts)} texts")

#         if patched_count:
#             logger.warning(
#                 f"⚠  {patched_count}/{len(texts)} chunks had NaN/Inf components "
#                 f"replaced with 0.0 — retrieval quality for those chunks may be "
#                 f"degraded. Check Ollama logs or investigate the input text."
#             )
#         if skipped_count:
#             logger.warning(
#                 f"⚠  {skipped_count}/{len(texts)} chunks could not be embedded and "
#                 f"will be missing from the index. Check Ollama logs for details."
#             )

#         return embeddings


# # ---------------------------------------------------------------------------
# # Chunking
# # ---------------------------------------------------------------------------

# class SectionAwareChunker:
#     """Advanced chunker that respects document section structure."""

#     def __init__(self, chunk_size: int = 512, overlap: int = 256):
#         self.chunk_size = chunk_size
#         self.overlap = overlap
#         self.section_patterns = [re.compile(p, re.MULTILINE) for p in SECTION_PATTERNS]

#     def detect_sections(self, text: str) -> List[DocumentSection]:
#         """Detect document sections and hierarchy from plain text."""
#         sections = []
#         lines = text.split('\n')
#         current_section: Dict = {"title": "Introduction", "content": "", "level": 0}
#         section_stack = [current_section]

#         for line in lines:
#             line_stripped = line.strip()

#             if not line_stripped:
#                 current_section["content"] += "\n"
#                 continue

#             is_header = False
#             header_level = 0

#             for pattern in self.section_patterns:
#                 match = pattern.match(line_stripped)
#                 if match:
#                     is_header = True
#                     if line_stripped.startswith('#'):
#                         header_level = len(line_stripped) - len(line_stripped.lstrip('#'))
#                     elif line_stripped.isupper():
#                         header_level = 1
#                     else:
#                         header_level = 2
#                     break

#             if is_header and len(line_stripped) < 200:
#                 if current_section["content"].strip():
#                     hierarchy = [s["title"] for s in section_stack if s["title"]]
#                     section_type = (
#                         "table"
#                         if current_section["title"].startswith("[Table ")
#                         else "text"
#                     )
#                     sections.append(DocumentSection(
#                         title=current_section["title"],
#                         content=current_section["content"].strip(),
#                         level=current_section["level"],
#                         section_type=section_type,
#                         section_hierarchy=list(hierarchy),
#                     ))

#                 while len(section_stack) > 1 and section_stack[-1]["level"] >= header_level:
#                     section_stack.pop()

#                 current_section = {
#                     "title": line_stripped.strip('#: ').strip(),
#                     "content": "",
#                     "level": header_level
#                 }
#                 section_stack.append(current_section)
#             else:
#                 current_section["content"] += line + "\n"

#         if current_section["content"].strip():
#             hierarchy = [s["title"] for s in section_stack if s["title"]]
#             section_type = (
#                 "table"
#                 if current_section["title"].startswith("[Table ")
#                 else "text"
#             )
#             sections.append(DocumentSection(
#                 title=current_section["title"],
#                 content=current_section["content"].strip(),
#                 level=current_section["level"],
#                 section_type=section_type,
#                 section_hierarchy=list(hierarchy),
#             ))

#         return sections if sections else [DocumentSection("Document", text, 0)]

#     def _split_long_sentence(self, sentence: str) -> List[str]:
#         """Split a sentence that exceeds MAX_CHUNK_SIZE at word boundaries."""
#         if len(sentence) <= MAX_CHUNK_SIZE:
#             return [sentence]
#         words = sentence.split()
#         parts: List[str] = []
#         current = ""
#         for word in words:
#             if len(current) + len(word) + 1 > MAX_CHUNK_SIZE and current:
#                 parts.append(current.strip())
#                 current = word
#             else:
#                 current = current + " " + word if current else word
#         if current.strip():
#             parts.append(current.strip())
#         return parts

#     def chunk_with_sentences(
#         self,
#         text: str,
#         section_title: str = "",
#         section_hierarchy: Optional[List[str]] = None
#     ) -> List[EnrichedChunk]:
#         """Chunk text respecting sentence boundaries with configurable overlap."""
#         if not text or len(text) < MIN_CHUNK_SIZE:
#             return []

#         if section_hierarchy is None:
#             section_hierarchy = [section_title]

#         raw_sentences = sent_tokenize(text)
#         sentences: List[str] = []
#         for s in raw_sentences:
#             sentences.extend(self._split_long_sentence(s))

#         chunks: List[EnrichedChunk] = []
#         current_chunk = ""
#         current_start = 0

#         for sentence in sentences:
#             if len(current_chunk) + len(sentence) > self.chunk_size and current_chunk:
#                 chunks.append(EnrichedChunk(
#                     text=current_chunk.strip(),
#                     section_title=section_title,
#                     section_hierarchy=list(section_hierarchy),
#                     page_number=None,
#                     chunk_type="text",
#                     word_count=len(word_tokenize(current_chunk)),
#                     sentence_count=len(sent_tokenize(current_chunk)),
#                     start_char=current_start,
#                     end_char=current_start + len(current_chunk)
#                 ))

#                 old_chunk_len = len(current_chunk)
#                 overlap_text = (
#                     current_chunk[-self.overlap:]
#                     if len(current_chunk) > self.overlap
#                     else current_chunk
#                 )
#                 current_chunk = overlap_text + " " + sentence
#                 current_start += old_chunk_len - len(overlap_text)
#             else:
#                 current_chunk += " " + sentence if current_chunk else sentence

#         if current_chunk.strip():
#             chunks.append(EnrichedChunk(
#                 text=current_chunk.strip(),
#                 section_title=section_title,
#                 section_hierarchy=list(section_hierarchy),
#                 page_number=None,
#                 chunk_type="text",
#                 word_count=len(word_tokenize(current_chunk)),
#                 sentence_count=len(sent_tokenize(current_chunk)),
#                 start_char=current_start,
#                 end_char=current_start + len(current_chunk)
#             ))

#         return chunks

#     def chunk_sections(self, sections: List[DocumentSection]) -> List[EnrichedChunk]:
#         """Chunk all document sections while preserving section context."""
#         all_chunks: List[EnrichedChunk] = []

#         for section in sections:
#             hierarchy = (
#                 section.section_hierarchy
#                 if hasattr(section, 'section_hierarchy')
#                 else [section.title]
#             )
#             section_chunks = self.chunk_with_sentences(section.content, section.title, hierarchy)

#             for chunk in section_chunks:
#                 chunk.page_number = section.page_number
#                 chunk.chunk_type = section.section_type

#             all_chunks.extend(section_chunks)

#         return all_chunks


# # ---------------------------------------------------------------------------
# # RAGAS HeadlineSplitter chunker (matches generate_dataset_v2.py exactly)
# # ---------------------------------------------------------------------------

# def _build_httpx_clients():
#     """
#     httpx clients tuned for multi-hour workloads against a local vLLM server.
#     Copied from generate_dataset_v2.py to keep behavior identical.
#     """
#     import httpx
#     limits = httpx.Limits(
#         max_keepalive_connections=0,
#         max_connections=100,
#         keepalive_expiry=0.0,
#     )
#     timeout = httpx.Timeout(300.0, connect=30.0)
#     sync_transport  = httpx.HTTPTransport(retries=3)
#     async_transport = httpx.AsyncHTTPTransport(retries=3)
#     return (
#         httpx.Client(limits=limits, timeout=timeout, transport=sync_transport),
#         httpx.AsyncClient(limits=limits, timeout=timeout, transport=async_transport),
#     )


# def _build_ragas_llm(provider: str, base_url: Optional[str] = None, model: Optional[str] = None):
#     """
#     Build a RAGAS-wrapped LLM for headline extraction.
#     Mirrors generate_dataset_v2.py -> build_llm() so dataset and index agree.
#     """
#     from ragas.llms import LangchainLLMWrapper

#     if provider == "nvidia":
#         from langchain_nvidia_ai_endpoints import ChatNVIDIA
#         if not os.environ.get("NVIDIA_API_KEY"):
#             raise RuntimeError("NVIDIA_API_KEY not set; required for --llm-provider nvidia.")
#         logger.info(f"  Headline LLM : NVIDIA NIM ({model or NVIDIA_LLM_MODEL})")
#         llm = ChatNVIDIA(
#             model=model or NVIDIA_LLM_MODEL,
#             nvidia_api_key=os.environ["NVIDIA_API_KEY"],
#             base_url=base_url or NVIDIA_BASE_URL,
#             temperature=0.1,
#             max_tokens=1024,
#         )
#     elif provider == "vllm":
#         from langchain_openai import ChatOpenAI
#         sync_client, async_client = _build_httpx_clients()
#         logger.info(f"  Headline LLM : vLLM ({model or VLLM_LLM_MODEL}) at {base_url or VLLM_BASE_URL}")
#         llm = ChatOpenAI(
#             model=model or VLLM_LLM_MODEL,
#             openai_api_key="dummy",
#             openai_api_base=base_url or VLLM_BASE_URL,
#             temperature=0.1,
#             max_tokens=1024,
#             http_client=sync_client,
#             http_async_client=async_client,
#         )
#     else:
#         raise ValueError(f"Unknown LLM provider: {provider!r}. Use 'nvidia' or 'vllm'.")

#     return LangchainLLMWrapper(llm)


# def _build_ragas_embeddings():
#     """
#     Build the embedding model RAGAS uses inside its KG transforms. This is
#     DIFFERENT from the embedder that writes to Qdrant — it only runs during
#     headline extraction / KG building. Using bge-m3 (same as dataset gen).
#     """
#     from ragas.embeddings import LangchainEmbeddingsWrapper
#     from langchain_huggingface import HuggingFaceEmbeddings

#     logger.info(f"  RAGAS internal embed : {RAGAS_INTERNAL_EMBED_MODEL}")
#     emb = HuggingFaceEmbeddings(
#         model_name=RAGAS_INTERNAL_EMBED_MODEL,
#         model_kwargs={"device": "cpu"},
#         encode_kwargs={"normalize_embeddings": True},
#     )
#     return LangchainEmbeddingsWrapper(emb)


# def _safe_headline_splitter():
#     """
#     HeadlineSplitter that silently skips nodes where no headlines were found
#     (same subclass used in generate_dataset_v2.py). Without this, pages with
#     no headings crash the entire pipeline.
#     """
#     from ragas.testset.transforms.splitters.headline import HeadlineSplitter

#     class SafeHeadlineSplitter(HeadlineSplitter):
#         async def split(self, node):
#             headlines = node.properties.get("headlines")
#             if not headlines:
#                 return [], []
#             return await super().split(node)

#     return SafeHeadlineSplitter()


# class RagasHeadlineChunker:
#     """
#     Alternative chunker that uses RAGAS's HeadlineSplitter — the SAME splitter
#     used by generate_dataset_v2.py to produce the dataset's reference_contexts.

#     Pipeline per PDF file:
#       1. LangChain PyPDFLoader → one Document per page
#       2. Each page → a NodeType.DOCUMENT in a fresh KnowledgeGraph
#       3. Run HeadlinesExtractor (LLM call per page) to label headlines
#       4. Run HeadlineSplitter to cut pages at those headlines → CHUNK nodes
#       5. Each CHUNK node's page_content becomes one EnrichedChunk

#     IMPORTANT: This makes an LLM call per page and is therefore much slower than
#     SectionAwareChunker. Do not use for first-pass ingestion of large corpora
#     unless you specifically need chunks that align with a RAGAS-generated eval
#     dataset.
#     """

#     def __init__(self, llm, embeddings, max_workers: int = 4):
#         self.llm = llm
#         self.embeddings = embeddings
#         self.max_workers = max_workers

#     def _run_config(self):
#         from ragas.run_config import RunConfig
#         return RunConfig(
#             timeout=600,
#             max_retries=15,
#             max_wait=180,
#             max_workers=self.max_workers,
#             seed=42,
#         )

#     def chunk_file(self, pdf_path: str) -> List[EnrichedChunk]:
#         """Run the full headline-extraction + split pipeline on a single PDF."""
#         from langchain_community.document_loaders import PyPDFLoader
#         from ragas.testset.graph import KnowledgeGraph, Node, NodeType
#         from ragas.testset.transforms import apply_transforms
#         from ragas.testset.transforms.extractors import HeadlinesExtractor

#         # 1. Load pages
#         loader = PyPDFLoader(pdf_path)
#         docs = loader.load()
#         if not docs:
#             logger.warning(f"  PyPDFLoader returned 0 pages for {pdf_path}")
#             return []

#         filename = os.path.basename(pdf_path)
#         for d in docs:
#             if "filename" not in d.metadata:
#                 d.metadata["filename"] = filename

#         # 2. Build tiny KG of DOCUMENT nodes
#         kg = KnowledgeGraph()
#         for d in docs:
#             kg.nodes.append(Node(
#                 type=NodeType.DOCUMENT,
#                 properties={
#                     "page_content": d.page_content,
#                     "document_metadata": d.metadata,
#                 },
#             ))

#         # 3. Minimal chunking pipeline — ONLY what's needed to produce chunks
#         # that align with generate_dataset_v2.py's reference_contexts.
#         #
#         # We explicitly skip the extras that default_transforms() includes
#         # (Summary, Themes, NER, CustomNodeFilter, Embedding, CosineSimilarity,
#         # OverlapScore) because those only exist to support TestsetGenerator's
#         # multi-hop question synthesis. For chunking alone they are dead LLM
#         # calls — ~7× slowdown with zero benefit.
#         trans = [
#             HeadlinesExtractor(llm=self.llm),
#             _safe_headline_splitter(),
#         ]

#         # 4. Apply transforms (LLM call per page for headline extraction, then split)
#         try:
#             apply_transforms(kg, trans, run_config=self._run_config())
#         except Exception as e:
#             logger.error(f"  ✗ RAGAS transforms failed for {filename}: {e}")
#             return []

#         # 5. Collect CHUNK nodes as EnrichedChunks
#         enriched: List[EnrichedChunk] = []
#         char_cursor = 0

#         for node in kg.nodes:
#             if node.type != NodeType.CHUNK:
#                 continue
#             content = node.properties.get("page_content", "") or ""
#             if not content.strip() or len(content) < MIN_CHUNK_SIZE:
#                 continue

#             # Headline (if any) goes into section_title for downstream metadata
#             headlines = node.properties.get("headlines") or []
#             if headlines:
#                 first_hl = headlines[0] if isinstance(headlines[0], str) else str(headlines[0])
#                 section_title = first_hl[:200]
#             else:
#                 section_title = "Document"

#             # Page number — best-effort: grab from parent document metadata if present
#             doc_meta = node.properties.get("document_metadata") or {}
#             page_num = doc_meta.get("page") if isinstance(doc_meta, dict) else None

#             enriched.append(EnrichedChunk(
#                 text=content.strip(),
#                 section_title=section_title,
#                 section_hierarchy=[section_title],
#                 page_number=page_num,
#                 chunk_type="text",
#                 word_count=len(word_tokenize(content)),
#                 sentence_count=len(sent_tokenize(content)),
#                 start_char=char_cursor,
#                 end_char=char_cursor + len(content),
#             ))
#             char_cursor += len(content)

#         logger.info(f"  ✓ RAGAS HeadlineSplitter produced {len(enriched)} chunks")
#         return enriched


# # ---------------------------------------------------------------------------
# # BM25 index
# # ---------------------------------------------------------------------------

# class BM25Index:
#     """BM25 sparse vector index."""

#     def __init__(self):
#         self.bm25 = None
#         self.tokenized_corpus: List[List[str]] = []
#         self.vocabulary: Dict[str, int] = {}
#         self.token_idf: Dict[str, float] = {}

#     def fit(self, texts: List[str]):
#         """Build BM25 index from a list of text strings."""
#         self.tokenized_corpus = [self._tokenize(t) for t in texts]
#         self.bm25 = BM25Okapi(self.tokenized_corpus)

#         all_tokens = sorted({t for doc in self.tokenized_corpus for t in doc})
#         self.vocabulary = {token: idx for idx, token in enumerate(all_tokens)}

#         N = len(self.tokenized_corpus)
#         for token in self.vocabulary:
#             df = sum(1 for doc in self.tokenized_corpus if token in doc)
#             self.token_idf[token] = np.log((N - df + 0.5) / (df + 0.5) + 1)

#         logger.info(f"  ✓ BM25 vocabulary size: {len(self.vocabulary)}")

#     def _tokenize(self, text: str) -> List[str]:
#         return [token.lower() for token in word_tokenize(text) if token.isalnum()]

#     def get_sparse_vector(self, text: str) -> SparseVector:
#         """Return a BM25 TF-IDF sparse vector for the given text."""
#         tokens = self._tokenize(text)
#         total = len(tokens)
#         token_counts: Dict[str, int] = {}

#         for token in tokens:
#             if token in self.vocabulary:
#                 token_counts[token] = token_counts.get(token, 0) + 1

#         indices: List[int] = []
#         values: List[float] = []

#         for token, count in token_counts.items():
#             tf = count / total if total else 0.0
#             idf = self.token_idf.get(token, 1.0)
#             indices.append(self.vocabulary[token])
#             values.append(float(tf * idf))

#         return SparseVector(indices=indices, values=values)


# # ---------------------------------------------------------------------------
# # Document loader (V2: pypdfium2 for PDF text, pdfplumber for table counts)
# # ---------------------------------------------------------------------------

# class AdvancedDocumentLoader:
#     """Enhanced document loader with pypdfium2-based PDF extraction."""

#     # TOC line pattern: text … dots … page number
#     _TOC_LINE_RE = re.compile(r'(\.\s*){2,}.*\d+\s*$')

#     @staticmethod
#     def _is_toc_page(page_text: str) -> bool:
#         """
#         Return True if this page looks like a Table-of-Contents / boilerplate
#         page that adds no retrievable content to the index.

#         Classification rules (all must pass to be considered TOC):

#         1. Near-empty pages (< TOC_MIN_CONTENT_CHARS chars) are always dropped —
#            these are blank separators or single-line chapter dividers.
#         2. Content-dense pages (> TOC_MAX_CONTENT_CHARS chars) are NEVER dropped
#            regardless of ratio, preventing false positives on real specification
#            pages that happen to contain a version annotation like
#            "Version ......... 3.0.0".
#         3. Pages with fewer than TOC_MIN_LINE_COUNT non-empty lines are NOT
#            ratio-classified — they are section headings, not TOC pages.
#         4. A page is classified as TOC only when MORE THAN TOC_LINE_RATIO of its
#            non-empty lines match the leader-dot pattern  (text . . . number).
#            The threshold is 0.50 (majority), not 0.30, to avoid false positives
#            from pages with a single dotted version line among real content.
#         """
#         stripped = page_text.strip()

#         # Rule 1 – near-empty page: blank separator, single-line chapter title
#         if len(stripped) < TOC_MIN_CONTENT_CHARS:
#             return True

#         # Rule 2 – content-dense page: too much text to be a pure TOC page
#         if len(stripped) > TOC_MAX_CONTENT_CHARS:
#             return False

#         # At this point stripped is between 50 and 800 chars of non-whitespace content.
#         # splitlines() on such a string always yields at least one non-empty line,
#         # so no explicit empty-list guard is needed here.
#         non_empty_lines = [l for l in stripped.splitlines() if l.strip()]

#         # Rule 3 – too few lines to make a reliable ratio decision
#         if len(non_empty_lines) < TOC_MIN_LINE_COUNT:
#             return False

#         # Rule 4 – majority of lines must be TOC-style entries
#         toc_line_count = sum(
#             1 for line in non_empty_lines
#             if AdvancedDocumentLoader._TOC_LINE_RE.search(line)
#         )
#         return (toc_line_count / len(non_empty_lines)) >= TOC_LINE_RATIO

#     @staticmethod
#     def extract_metadata(path: str) -> Dict:
#         """Extract filesystem metadata."""
#         file_stat = os.stat(path)
#         return {
#             "file_size_bytes": file_stat.st_size,
#             "created_timestamp": file_stat.st_ctime,
#             "modified_timestamp": file_stat.st_mtime,
#             "file_extension": Path(path).suffix.lower(),
#         }

#     @staticmethod
#     def load_pdf(path: str) -> Tuple[str, Dict]:
#         """
#         Load PDF text using pypdfium2 (proper word spacing) and gather table-count
#         metadata using pdfplumber (table extraction only, not for text content).

#         TOC pages and near-empty pages are skipped automatically.
#         All PDF handles are closed in finally blocks to prevent file-descriptor
#         leaks even when a page is corrupted or raises an exception mid-loop.
#         """
#         metadata: Dict = {"num_pages": 0, "has_tables": False, "tables_count": 0}
#         text_parts: List[str] = []
#         skipped_pages: List[int] = []

#         # --- text extraction via pypdfium2 ---
#         pdf = pdfium.PdfDocument(path)
#         try:
#             num_pages = len(pdf)
#             metadata["num_pages"] = num_pages

#             for page_num in range(num_pages):
#                 page = pdf[page_num]
#                 page_text = ""          # initialise so the variable is always bound
#                 try:
#                     textpage = page.get_textpage()
#                     try:
#                         page_text = textpage.get_text_bounded()
#                     finally:
#                         textpage.close()
#                 except Exception as page_err:
#                     # Corrupted or unreadable page — skip it and continue with the rest
#                     logger.warning(
#                         f"  ⚠ Could not extract text from page {page_num + 1}: {page_err}"
#                     )
#                 finally:
#                     page.close()

#                 if not page_text:
#                     continue

#                 if AdvancedDocumentLoader._is_toc_page(page_text):
#                     skipped_pages.append(page_num + 1)
#                     continue

#                 text_parts.append(f"\n[Page {page_num + 1}]\n{page_text}\n")
#         finally:
#             pdf.close()

#         if skipped_pages:
#             logger.info(
#                 f"  ⊘ Skipped {len(skipped_pages)} TOC/boilerplate pages: "
#                 f"{skipped_pages[:10]}{'...' if len(skipped_pages) > 10 else ''}"
#             )

#         # --- table count metadata via pdfplumber (no text used) ---
#         try:
#             with pdfplumber.open(path) as plumber_pdf:
#                 for page in plumber_pdf.pages:
#                     tables = page.extract_tables()
#                     if tables:
#                         metadata["has_tables"] = True
#                         metadata["tables_count"] += len(tables)
#         except Exception as e:
#             logger.warning(f"  pdfplumber table-count failed (non-fatal): {e}")

#         return "".join(text_parts), metadata

#     @staticmethod
#     def load_docx(path: str) -> Tuple[str, Dict]:
#         """Load DOCX with metadata."""
#         doc = docx.Document(path)
#         text_parts: List[str] = []
#         metadata: Dict = {"num_paragraphs": 0, "has_tables": False, "tables_count": 0}

#         for paragraph in doc.paragraphs:
#             if paragraph.text.strip():
#                 text_parts.append(paragraph.text)
#                 metadata["num_paragraphs"] += 1

#         if doc.tables:
#             metadata["has_tables"] = True
#             metadata["tables_count"] = len(doc.tables)
#             for table in doc.tables:
#                 for row in table.rows:
#                     row_text = " | ".join([cell.text for cell in row.cells])
#                     text_parts.append(row_text)

#         return "\n".join(text_parts), metadata

#     @staticmethod
#     def load_txt(path: str) -> Tuple[str, Dict]:
#         """Load TXT file."""
#         with open(path, "r", encoding="utf-8", errors="ignore") as f:
#             text = f.read()
#         metadata = {
#             "num_lines": len(text.split('\n')),
#             "char_count": len(text)
#         }
#         return text, metadata

#     @classmethod
#     def load(cls, path: str) -> Tuple[Optional[str], Dict]:
#         """Universal loader: dispatches to the right format handler."""
#         try:
#             base_metadata = cls.extract_metadata(path)

#             if path.lower().endswith(".pdf"):
#                 text, doc_metadata = cls.load_pdf(path)
#             elif path.lower().endswith(".docx"):
#                 text, doc_metadata = cls.load_docx(path)
#             elif path.lower().endswith(".txt"):
#                 text, doc_metadata = cls.load_txt(path)
#             else:
#                 return None, {}

#             base_metadata.update(doc_metadata)
#             return text, base_metadata

#         except Exception as e:
#             logger.error(f"Error loading {path}: {e}")
#             return None, {}


# # ---------------------------------------------------------------------------
# # Main ingestion pipeline
# # ---------------------------------------------------------------------------

# def main():
#     parser = argparse.ArgumentParser(description="Advanced RAG Ingestion System V2")
#     parser.add_argument('--data-dir', type=str, default=DATA_DIR,
#                         help='Directory containing documents')
#     parser.add_argument('--collection', type=str, default=COLLECTION,
#                         help='Qdrant collection name')
#     parser.add_argument('--qdrant-url', type=str, default=QDRANT_URL,
#                         help='Qdrant server URL')
#     parser.add_argument('--ollama-url', type=str, default=OLLAMA_URL,
#                         help='Ollama server URL')
#     parser.add_argument('--ollama-model', type=str, default=OLLAMA_MODEL,
#                         help='Ollama embedding model name')
#     parser.add_argument('--chunk-size', type=int, default=CHUNK_SIZE,
#                         help='Chunk size in characters')
#     parser.add_argument('--chunk-overlap', type=int, default=CHUNK_OVERLAP,
#                         help='Chunk overlap in characters')
#     parser.add_argument('--bm25-output', type=str, default=BM25_OUTPUT,
#                         help='Path to save BM25 index JSON')
#     parser.add_argument('--chunker', type=str, choices=['section', 'ragas'],
#                         default='section',
#                         help="Chunker to use. 'section' (default) = existing "
#                              "SectionAwareChunker. 'ragas' = RAGAS HeadlineSplitter "
#                              "(matches generate_dataset_v2.py — REQUIRED if you want "
#                              "chunks to align with a RAGAS-generated eval dataset). "
#                              "'ragas' makes an LLM call per page and is much slower.")
#     parser.add_argument('--llm-provider', type=str, choices=['nvidia', 'vllm'],
#                         default='nvidia',
#                         help="LLM provider for HeadlinesExtractor when --chunker=ragas. "
#                              "Ignored otherwise.")
#     parser.add_argument('--llm-base-url', type=str, default=None,
#                         help="Override the LLM base URL. Defaults to NVIDIA_BASE_URL "
#                              "or VLLM_BASE_URL depending on --llm-provider.")
#     parser.add_argument('--llm-model', type=str, default=None,
#                         help="Override the LLM model name for HeadlinesExtractor.")
#     parser.add_argument('--ragas-max-workers', type=int, default=4,
#                         help="Parallel workers for RAGAS transforms (default 4). "
#                              "Only used when --chunker=ragas.")
#     args = parser.parse_args()

#     data_dir = args.data_dir
#     collection = args.collection
#     qdrant_url = args.qdrant_url
#     ollama_url = args.ollama_url
#     ollama_model = args.ollama_model
#     chunk_size = args.chunk_size
#     chunk_overlap = args.chunk_overlap
#     bm25_output = args.bm25_output

#     logger.info("=" * 80)
#     logger.info("ADVANCED RAG INGESTION SYSTEM V2")
#     logger.info("=" * 80)

#     # --- Qdrant client ---
#     client = QdrantClient(url=qdrant_url)

#     # --- Embedding model ---
#     if USE_OLLAMA_BGE_M3:
#         try:
#             logger.info(f"Using Ollama BGE-M3 model at {ollama_url}")
#             embedder = OllamaBGEM3Embedder(ollama_url, ollama_model)
#             embedding_dim = embedder.dimension
#         except Exception:
#             logger.warning(f"Ollama unavailable, falling back to {FALLBACK_MODEL}")
#             embedder = SentenceTransformer(FALLBACK_MODEL)
#             embedding_dim = 384
#     else:
#         logger.info(f"Using SentenceTransformer: {FALLBACK_MODEL}")
#         embedder = SentenceTransformer(FALLBACK_MODEL)
#         embedding_dim = 384

#     # --- Collection setup ---
#     existing = [c.name for c in client.get_collections().collections]
#     if collection not in existing:
#         client.create_collection(
#             collection_name=collection,
#             vectors_config={
#                 "dense": VectorParams(size=embedding_dim, distance=Distance.COSINE)
#             },
#             sparse_vectors_config={
#                 "bm25": SparseVectorParams(index=SparseIndexParams(on_disk=False))
#             },
#         )
#         logger.info(f"✓ Created collection: {collection}")
#     else:
#         col_info = client.get_collection(collection)
#         existing_dim = None
#         if hasattr(col_info.config.params.vectors, "get"):
#             dense_cfg = col_info.config.params.vectors.get("dense")
#             if dense_cfg is not None:
#                 existing_dim = dense_cfg.size
#         if existing_dim is not None and existing_dim != embedding_dim:
#             logger.error(
#                 f"✗ Dimension mismatch: collection '{collection}' has {existing_dim}-dim vectors, "
#                 f"but current embedder produces {embedding_dim}-dim vectors. Aborting."
#             )
#             return
#         logger.info(f"✓ Collection exists: {collection}")

#     # --- Pipeline components ---
#     if args.chunker == "ragas":
#         logger.info("\nInitializing RAGAS HeadlineSplitter chunker (LLM-backed)...")
#         ragas_llm = _build_ragas_llm(
#             provider=args.llm_provider,
#             base_url=args.llm_base_url,
#             model=args.llm_model,
#         )
#         # Note: minimal pipeline (HeadlinesExtractor + HeadlineSplitter) does
#         # not need an embedding model — we skip building one for faster startup.
#         chunker = RagasHeadlineChunker(
#             llm=ragas_llm,
#             embeddings=None,
#             max_workers=args.ragas_max_workers,
#         )
#     else:
#         chunker = SectionAwareChunker(chunk_size, chunk_overlap)
#     bm25_index = BM25Index()

#     total_chunks = 0
#     total_files = 0
#     all_points: List[Tuple[Dict, str]] = []

#     logger.info("\nPhase 1: Loading and chunking documents...")

#     for root, _, files in os.walk(data_dir):
#         for file in files:
#             if not file.lower().endswith((".pdf", ".docx", ".txt")):
#                 continue

#             path = os.path.join(root, file)
#             logger.info(f"\n→ Processing: {file}")

#             try:
#                 h = file_hash(path)

#                 if already_indexed(client, collection, h):
#                     logger.info("  ⊘ Already indexed")
#                     continue

#                 text, doc_metadata = AdvancedDocumentLoader.load(path)
#                 if not text or len(text.strip()) < MIN_CHUNK_SIZE:
#                     logger.info("  ⊘ Empty or too short after extraction")
#                     continue

#                 # --- Chunking ---
#                 if args.chunker == "ragas":
#                     # RAGAS chunker re-loads the PDF via PyPDFLoader internally
#                     # (identical to generate_dataset_v2.py) so chunks match the
#                     # dataset exactly. Only valid for PDFs.
#                     if not path.lower().endswith(".pdf"):
#                         logger.info(f"  ⊘ --chunker=ragas only supports PDFs; skipping {file}")
#                         continue
#                     raw_chunks = chunker.chunk_file(path)
#                     logger.info(f"  ✓ RAGAS chunker produced {len(raw_chunks)} chunks")
#                 elif ENABLE_SECTION_AWARE:
#                     sections = chunker.detect_sections(text)
#                     raw_chunks = chunker.chunk_sections(sections)
#                     logger.info(
#                         f"  ✓ Detected {len(sections)} sections → {len(raw_chunks)} raw chunks"
#                     )
#                 else:
#                     raw_chunks = chunker.chunk_with_sentences(text, "Document")
#                     logger.info(f"  ✓ Created {len(raw_chunks)} raw chunks")

#                 if not raw_chunks:
#                     continue

#                 # --- Near-duplicate chunk removal ---
#                 accepted_chunks = []
#                 dedup_window: List[str] = []   # sliding window of last JACCARD_WINDOW texts

#                 for chunk in raw_chunks:
#                     is_dup = False
#                     for prev_text in dedup_window[-JACCARD_WINDOW:]:
#                         if jaccard_similarity(chunk.text, prev_text) > JACCARD_DEDUP_THRESHOLD:
#                             is_dup = True
#                             break
#                     if not is_dup:
#                         accepted_chunks.append(chunk)
#                         dedup_window.append(chunk.text)

#                 dedup_removed = len(raw_chunks) - len(accepted_chunks)
#                 if dedup_removed:
#                     logger.info(f"  ⊘ Removed {dedup_removed} near-duplicate chunks")

#                 chunks = accepted_chunks
#                 if not chunks:
#                     continue

#                 chunk_texts = [chunk.text for chunk in chunks]

#                 # --- Embeddings ---
#                 logger.info(f"  ⚡ Generating embeddings for {len(chunk_texts)} chunks...")
#                 if isinstance(embedder, OllamaBGEM3Embedder):
#                     embeddings = embedder.encode(chunk_texts, batch_size=8)
#                 else:
#                     raw_embeddings = embedder.encode(
#                         chunk_texts,
#                         batch_size=32,
#                         show_progress_bar=False,
#                         convert_to_numpy=True
#                     )
#                     embeddings = [emb.tolist() for emb in raw_embeddings]

#                 # --- Build Qdrant point records ---
#                 for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
#                     if embedding is None:
#                         logger.warning(f"  ⚠ Skipping chunk {i} — embedding failed")
#                         continue
#                     point_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, f"{h}_{i}"))
#                     point = {
#                         "id": point_id,
#                         "vector": embedding,
#                         "payload": {
#                             "content": chunk.text,
#                             "source_path": path,
#                             "filename": file,
#                             "folder": os.path.relpath(root, data_dir),
#                             "file_type": Path(file).suffix.lower(),
#                             "file_hash": h,
#                             "chunk_id": i,
#                             "total_chunks": len(chunks),
#                             "chunk_length": len(chunk.text),
#                             "word_count": chunk.word_count,
#                             "sentence_count": chunk.sentence_count,
#                             "start_char": chunk.start_char,
#                             "end_char": chunk.end_char,
#                             "section_title": chunk.section_title,
#                             "section_hierarchy": chunk.section_hierarchy,
#                             "chunk_type": chunk.chunk_type,
#                             **doc_metadata
#                         }
#                     }
#                     all_points.append((point, chunk.text))
#                     total_chunks += 1

#                 total_files += 1

#             except Exception as e:
#                 logger.error(f"  ✗ Error processing {file}: {e}")
#                 continue

#     if not all_points:
#         logger.warning("No documents to index!")
#         return

#     # --- Phase 2: BM25 ---
#     all_texts_for_bm25 = [text for _, text in all_points]
#     logger.info(f"\nPhase 2: Building BM25 index for {len(all_texts_for_bm25)} chunks...")
#     bm25_index.fit(all_texts_for_bm25)

#     bm25_data = {
#         "vocabulary": bm25_index.vocabulary,
#         "token_idf": {k: float(v) for k, v in bm25_index.token_idf.items()},
#     }
#     Path(bm25_output).parent.mkdir(parents=True, exist_ok=True)
#     with open(bm25_output, "w") as f:
#         json.dump(bm25_data, f)
#     logger.info(f"  ✓ Saved BM25 index to {bm25_output}")

#     # --- Phase 3: Upload to Qdrant ---
#     logger.info("\nPhase 3: Adding sparse vectors and uploading to Qdrant...")

#     points_to_upload: List[PointStruct] = []
#     for point_data, chunk_text in all_points:
#         sparse_vector = bm25_index.get_sparse_vector(chunk_text)
#         points_to_upload.append(PointStruct(
#             id=point_data["id"],
#             vector={
#                 "dense": point_data["vector"],
#                 "bm25": sparse_vector
#             },
#             payload=point_data["payload"]
#         ))

#     batch_size = 100
#     num_batches = (len(points_to_upload) - 1) // batch_size + 1
#     for i in range(0, len(points_to_upload), batch_size):
#         batch = points_to_upload[i:i + batch_size]
#         client.upsert(collection_name=collection, points=batch, wait=True)
#         logger.info(f"  ✓ Uploaded batch {i // batch_size + 1}/{num_batches}")

#     # --- Summary ---
#     logger.info("\n" + "=" * 80)
#     logger.info("INGESTION COMPLETE")
#     logger.info("=" * 80)
#     logger.info(f"✓ Files processed:      {total_files}")
#     logger.info(f"✓ Total chunks indexed: {total_chunks}")
#     logger.info(
#         f"✓ Avg chunks/file:     "
#         f"{total_chunks / total_files if total_files > 0 else 0:.1f}"
#     )
#     logger.info(f"✓ Collection:           {collection}")
#     logger.info(f"✓ Chunker used:         {args.chunker}")
#     logger.info(f"✓ Embedding dimension:  {embedding_dim}")
#     logger.info(f"✓ BM25 vocabulary size: {len(bm25_index.vocabulary)}")
#     logger.info("=" * 80)


# if __name__ == "__main__":
#     main()




"""
ADVANCED RAG INGESTION SYSTEM V2 (FIXED)
=========================================
Fixes applied over the previous version:

FIX 1 — MAX_EMBED_CHARS raised from 512 → 8000
    BGE-M3 has an 8192-token context window (~32 000 chars).  The old value of
    512 silently truncated every chunk larger than 512 chars before embedding,
    making the second half of 90% of all chunks invisible to dense search.
    8000 chars (~2000 tokens) is well within BGE-M3's window and safely covers
    any chunk produced by this pipeline.

FIX 2 — SectionAwareChunker AUTOSAR patterns added
    All four original regex patterns failed to match any real AUTOSAR section
    header ("Test Objective", "Main Test Execution", "4.3.5 [ATS_…]", etc.).
    New patterns cover: AUTOSAR numbered sections with bracket IDs, standard
    field labels (Test Objective, Pre-conditions, Main Test Execution, etc.),
    and requirement identifiers.  The original patterns are preserved so the
    chunker still works on Markdown / generic docs.

    NOTE: For maximum alignment with a RAGAS-generated evaluation dataset,
    always run ingestion with --chunker=ragas.  The SectionAwareChunker is
    provided as a fast fallback for future document additions.

FIX 3 — BM25 tokenizer preserves AUTOSAR compound identifiers
    The original tokenizer applied word_tokenize + isalnum, which split
    "SWS_Com_00228" → ['sws','com','00228'] and destroyed identifier specificity.
    The new tokenizer uses a two-pass strategy:
      Pass 1: extract whole compound tokens (containing _ or :: or digits
              joined to letters) and emit them as-is (lowercased).
      Pass 2: also emit the individual alphanumeric parts for partial matching.
    The SAME tokenizer is used for both document indexing and query encoding
    (see also the matching fix in Evaluate_Retrieval_With_Reranker_Template.py).

All original V2 features are preserved: pypdfium2 extraction, TOC detection,
near-duplicate removal, Jaccard dedup, embedding retry logic, deduplication
by file hash, dimension-mismatch check, RAGAS HeadlineSplitter chunker.
"""

import argparse
import json
import math
import numpy as np
import os
import re
import hashlib
import uuid
import time
from typing import Optional, List, Dict, Tuple
from dataclasses import dataclass
from pathlib import Path
import logging

import pypdfium2 as pdfium
import pdfplumber
import docx
import requests
from sentence_transformers import SentenceTransformer
from qdrant_client import QdrantClient
from qdrant_client.models import (
    VectorParams,
    Distance,
    Filter,
    FieldCondition,
    MatchValue,
    PointStruct,
    SparseVector,
    SparseVectorParams,
    SparseIndexParams,
)
from rank_bm25 import BM25Okapi
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize

try:
    nltk.data.find('tokenizers/punkt')
except LookupError:
    try:
        nltk.download('punkt_tab', quiet=True)
    except Exception:
        nltk.download('punkt', quiet=True)

# ================= CONFIG =================

DATA_DIR    = r"/home/olj3kor/praveen/Autosar_docs_2"
COLLECTION  = "Autosar_v2"
QDRANT_URL  = "http://localhost:7333"

USE_OLLAMA_BGE_M3 = True
OLLAMA_URL        = "http://localhost:11434"
OLLAMA_MODEL      = "bge-m3:latest"
FALLBACK_MODEL    = "sentence-transformers/all-MiniLM-L6-v2"

CHUNK_SIZE     = 1024
CHUNK_OVERLAP  = 256
MIN_CHUNK_SIZE = 100
MAX_CHUNK_SIZE = 1024

ENABLE_SECTION_AWARE = True

# ── FIX 2: AUTOSAR-aware section patterns ─────────────────────────────────
# Original generic patterns kept first so Markdown / plain-text docs still work.
# AUTOSAR-specific patterns added after.
SECTION_PATTERNS = [
    # --- original generic patterns ---
    r'^#{1,6}\s+(.+)$',                        # Markdown headers
    r'^([A-Z][^.!?]*):$',                       # Title case with colon
    r'^\d+\.\s+([A-Z].+)$',                    # Numbered sections (generic)
    r'^([A-Z\s]{3,})$',                         # All-caps headers

    # --- AUTOSAR-specific patterns (FIX 2) ---
    # Numbered sections with bracket requirement IDs:
    #   "4.3.5 [ATS_COMFR_00231] Signal on Time Base…"
    #   "7.1 [SWS_Com_00228] Transmission Mode"
    r'^\d+(?:\.\d+)*\s+\[[A-Z_0-9]+\].*$',

    # Numbered sections without brackets but with mixed-case title:
    #   "4.3.6 Signal Group on Time Base"
    r'^\d+(?:\.\d+)+\s+[A-Z].*$',

    # Standard AUTOSAR test-case field labels (exact match, any case):
    r'^(?:Test Objective|Test Steps?|Pass Criteria|Pre-?conditions?|'
    r'Post-?conditions?|Main Test Execution|Configuration Parameters?|'
    r'Needed Adaptation to other Releases?|Summary|Trace to Requirement|'
    r'Affected Modules?|State|Test System|Test Configuration|'
    r'Test Case Design|Re-?usable Test Steps?|Rationale|'
    r'Use Case|Applicability to Car Domains?)[\s:]*$',

    # Requirement / specification identifiers used as section headers:
    #   "[SWS_Com_00228]"  "[PRS_SOMEIP_00201]"
    r'^\[[A-Z]+_[A-Z0-9_]+\]\s*.*$',

    # All-uppercase AUTOSAR acronym headers:
    #   "AUTOSAR"  "COM"  "VFB"
    r'^[A-Z][A-Z0-9_\s]{2,}$',
]

BM25_OUTPUT = "ingestion_output/bm25_index.json"

JACCARD_DEDUP_THRESHOLD = 0.90
JACCARD_WINDOW          = 20

TOC_LINE_RATIO        = 0.50
TOC_MIN_CONTENT_CHARS = 50
TOC_MIN_LINE_COUNT    = 5
TOC_MAX_CONTENT_CHARS = 800

# ── FIX 1: raised from 512 to 8000 ───────────────────────────────────────
# BGE-M3 context window = 8192 tokens ≈ 32 000 chars.
# Old value (512 chars) silently truncated 90 % of all chunks, making the
# second half of every large chunk invisible to dense retrieval.
MAX_EMBED_CHARS = 8000

NVIDIA_LLM_MODEL           = "meta/llama-3.1-70b-instruct"
NVIDIA_BASE_URL            = "https://integrate.api.nvidia.com/v1"
VLLM_LLM_MODEL             = "Qwen/Qwen2.5-72B-Instruct-AWQ"
VLLM_BASE_URL              = "http://localhost:8011/v1"
RAGAS_INTERNAL_EMBED_MODEL = "BAAI/bge-m3"

# ── Ingestion audit log ────────────────────────────────────────────────────
# Every ingestion run writes two files into this directory:
#   ingestion_<timestamp>.log  — full console output (mirrors what you see on screen)
#   ingestion_<timestamp>.json — structured per-file and per-chunk audit trail
# The BM25 index (bm25_index.json) is also written here.
# All three outputs stay together in one place for easy inspection.
INGESTION_OUTPUT_DIR = "./ingestion_output"
INGESTION_LOG_DIR    = INGESTION_OUTPUT_DIR

# =========================================

import sys

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


def _setup_file_logging(log_dir: str) -> Tuple[str, str]:
    """
    Creates two output files for this run and wires a file handler into the
    root logger so every logger.info/warning/error call goes to both the
    console AND the log file simultaneously.

    Returns (log_file_path, json_file_path).
    """
    Path(log_dir).mkdir(parents=True, exist_ok=True)
    ts = time.strftime("%Y%m%d_%H%M%S")
    log_path  = str(Path(log_dir) / f"ingestion_{ts}.log")
    json_path = str(Path(log_dir) / f"ingestion_{ts}.json")

    file_handler = logging.FileHandler(log_path, encoding="utf-8")
    file_handler.setLevel(logging.DEBUG)          # capture DEBUG too in the file
    file_handler.setFormatter(logging.Formatter('%(asctime)s - %(levelname)s - %(message)s'))
    logging.getLogger().addHandler(file_handler)

    logger.info(f"Ingestion log  : {log_path}")
    logger.info(f"Audit JSON     : {json_path}")
    return log_path, json_path


class IngestionAudit:
    """
    Collects every observable fact produced during ingestion and writes them
    to a structured JSON file so you can diagnose exactly what happened for
    every file, every chunk, and every embedding call without re-running.

    Structure of the output JSON:
    {
      "run": { timestamp, config, chunker, embedding_model, ... },
      "files": [
        {
          "filename": "AUTOSAR_ATS_CommunicationViaBus.pdf",
          "status": "ok" | "skipped_already_indexed" | "skipped_empty" | "error",
          "file_hash": "...",
          "num_pages": 165,
          "pages_skipped_toc": [1, 2, 3],
          "raw_chunks": 48,
          "dedup_removed": 2,
          "final_chunks": 46,
          "embedding_skipped": 0,
          "embedding_patched": 0,
          "chunks": [
            {
              "chunk_id": 0,
              "section_title": "Test Objective",
              "page_number": 115,
              "char_len": 1271,
              "word_count": 210,
              "chunk_type": "text",
              "text_preview": "Main Test Execution  Test Steps...",
              "headline_extraction": {        <- RAGAS only
                "headlines_found": ["Main Test Execution"],
                "source_page": 115
              },
              "embedding_status": "ok" | "skipped" | "patched",
              "bm25_token_count": 87,
              "bm25_compound_tokens": ["at_231_ipdugroup", "sws_com_00228"]
            }, ...
          ]
        }, ...
      ],
      "bm25": {
        "vocabulary_size": 25016,
        "compound_token_count": 6368,
        "sample_compound_tokens": [...],
        "critical_tokens_present": { "sws_com_00228": true, "prs_e2e_00085": false, ... },
        "idf_stats": { "mean": 7.13, "max": 7.99, "min": 0.097 }
      },
      "upload": {
        "total_points": 4426,
        "batches": 45,
        "errors": []
      },
      "summary": {
        "total_files": 12,
        "total_chunks": 4426,
        "total_embedding_skipped": 3,
        "total_embedding_patched": 1,
        "files_with_zero_chunks": [],
        "files_with_all_document_sections": [],
        "warnings": []
      }
    }
    """

    # Tokens we always check in the BM25 vocab — the ones that caused the
    # prs_e2e_00085 regression.  Extend this list as you add new doc families.
    CRITICAL_TOKENS = [
        "sws_com_00228", "sws_com_00768", "sws_com_00495",
        "at_231_ipdugroup", "at_275_ipdugroup", "at_231_sg1",
        "prs_e2e_00085", "prs_e2e_00086",
        "e2e_p01dataidmode", "e2e_p02dataidmode",
        "prs_someip_00201", "prs_someip_00228", "prs_someip_00230",
        "ats_comfr_00231", "ats_comfr_00232",
        "ara::com", "bswmmoderequest", "ipdu_activated", "rte_switch",
    ]

    def __init__(self, json_path: str, run_config: Dict):
        self.json_path = json_path
        self.data: Dict = {
            "run":     run_config,
            "files":   [],
            "bm25":    {},
            "upload":  {"total_points": 0, "batches": 0, "errors": []},
            "summary": {
                "total_files": 0,
                "total_chunks": 0,
                "total_embedding_skipped": 0,
                "total_embedding_patched": 0,
                "files_with_zero_chunks": [],
                "files_with_all_document_sections": [],
                "warnings": [],
            },
        }
        self._current_file: Optional[Dict] = None

    # ── File lifecycle ─────────────────────────────────────────────────────

    def start_file(self, filename: str, file_hash: str, path: str):
        self._current_file = {
            "filename":            filename,
            "path":                path,
            "file_hash":           file_hash,
            "status":              "processing",
            "num_pages":           None,
            "pages_skipped_toc":   [],
            "raw_chunks":          0,
            "dedup_removed":       0,
            "final_chunks":        0,
            "embedding_skipped":   0,
            "embedding_patched":   0,
            "chunks":              [],
            "error":               None,
        }

    def set_file_pages(self, num_pages: int, skipped_toc: List[int]):
        if self._current_file:
            self._current_file["num_pages"]         = num_pages
            self._current_file["pages_skipped_toc"] = skipped_toc

    def set_file_chunk_counts(self, raw: int, dedup_removed: int, final: int):
        if self._current_file:
            self._current_file["raw_chunks"]    = raw
            self._current_file["dedup_removed"] = dedup_removed
            self._current_file["final_chunks"]  = final

    def set_file_embedding_counts(self, skipped: int, patched: int):
        if self._current_file:
            self._current_file["embedding_skipped"] = skipped
            self._current_file["embedding_patched"] = patched
            self.data["summary"]["total_embedding_skipped"] += skipped
            self.data["summary"]["total_embedding_patched"] += patched

    def finish_file(self, status: str = "ok", error: str = None):
        if self._current_file:
            self._current_file["status"] = status
            if error:
                self._current_file["error"] = error
            # Detect all-'Document' section titles — the silent RAGAS failure signal
            if status == "ok" and self._current_file["chunks"]:
                titles = [c["section_title"] for c in self._current_file["chunks"]]
                unique = set(titles)
                if unique == {"Document"}:
                    fname = self._current_file["filename"]
                    self.data["summary"]["files_with_all_document_sections"].append(fname)
                    self.data["summary"]["warnings"].append(
                        f"{fname}: ALL {len(titles)} chunks have section_title='Document' "
                        f"— HeadlinesExtractor likely returned empty for every page."
                    )
                    logger.warning(
                        f"  ⚠ ALL {len(titles)} chunks for {fname} have "
                        f"section_title='Document'. HeadlinesExtractor may have failed "
                        f"for every page in this file. Check the LLM endpoint and logs."
                    )
            if status == "ok":
                self.data["summary"]["total_files"]  += 1
                self.data["summary"]["total_chunks"] += self._current_file["final_chunks"]
            elif status == "skipped_zero_chunks":
                self.data["summary"]["files_with_zero_chunks"].append(
                    self._current_file["filename"]
                )
            self.data["files"].append(self._current_file)
            self._current_file = None

    def skip_file(self, filename: str, path: str, reason: str):
        self.data["files"].append({
            "filename": filename,
            "path":     path,
            "status":   reason,
            "chunks":   [],
        })

    # ── Chunk recording ────────────────────────────────────────────────────

    def record_chunk(
        self,
        chunk_id:         int,
        chunk:            "EnrichedChunk",
        embedding_status: str,
        bm25_tokens:      List[str],
        headline_info:    Optional[Dict] = None,
    ):
        if not self._current_file:
            return
        compound_tokens = [t for t in bm25_tokens if "_" in t or "::" in t]
        self._current_file["chunks"].append({
            "chunk_id":           chunk_id,
            "section_title":      chunk.section_title,
            "section_hierarchy":  chunk.section_hierarchy,
            "page_number":        chunk.page_number,
            "char_len":           len(chunk.text),
            "word_count":         chunk.word_count,
            "chunk_type":         chunk.chunk_type,
            "text_preview":       chunk.text[:120].replace("\n", " "),
            "text_tail":          chunk.text[-80:].replace("\n", " "),
            "embedding_status":   embedding_status,
            "bm25_token_count":   len(bm25_tokens),
            "bm25_compound_tokens": compound_tokens[:20],  # first 20 compound tokens
            **({"headline_info": headline_info} if headline_info else {}),
        })

    # ── BM25 audit ─────────────────────────────────────────────────────────

    def record_bm25(self, bm25_index: "BM25Index"):
        import statistics as _stats
        vocab      = bm25_index.vocabulary
        token_idf  = bm25_index.token_idf
        compound   = [t for t in vocab if "_" in t or "::" in t]
        all_idfs   = list(token_idf.values())

        critical_check = {
            t: (t in vocab) for t in self.CRITICAL_TOKENS
        }
        missing_critical = [t for t, present in critical_check.items() if not present]
        if missing_critical:
            msg = f"BM25: {len(missing_critical)} critical tokens MISSING from vocab: {missing_critical}"
            self.data["summary"]["warnings"].append(msg)
            logger.warning(f"  ⚠ {msg}")
        else:
            logger.info(f"  ✓ BM25: all {len(self.CRITICAL_TOKENS)} critical tokens present in vocab")

        self.data["bm25"] = {
            "vocabulary_size":          len(vocab),
            "compound_token_count":     len(compound),
            "compound_token_pct":       round(100 * len(compound) / max(len(vocab), 1), 1),
            "sample_compound_tokens":   sorted(compound)[:50],
            "critical_tokens_present":  critical_check,
            "missing_critical_tokens":  missing_critical,
            "idf_stats": {
                "mean":              round(_stats.mean(all_idfs), 4)   if all_idfs else 0,
                "median":            round(_stats.median(all_idfs), 4) if all_idfs else 0,
                "max":               round(max(all_idfs), 4)           if all_idfs else 0,
                "min":               round(min(all_idfs), 4)           if all_idfs else 0,
                "tokens_idf_gte_7":  sum(1 for v in all_idfs if v >= 7.0),
                "tokens_idf_lte_2":  sum(1 for v in all_idfs if v <= 2.0),
            },
        }

    # ── Upload audit ───────────────────────────────────────────────────────

    def record_upload_batch(self, batch_num: int, batch_size: int, error: str = None):
        self.data["upload"]["batches"]       += 1
        self.data["upload"]["total_points"]  += batch_size
        if error:
            self.data["upload"]["errors"].append({"batch": batch_num, "error": error})

    # ── Flush to disk ──────────────────────────────────────────────────────

    def save(self):
        """Write the full audit JSON. Called after every major phase."""
        try:
            with open(self.json_path, "w", encoding="utf-8") as f:
                json.dump(self.data, f, indent=2, default=str)
        except Exception as e:
            logger.error(f"Could not write audit JSON to {self.json_path}: {e}")

    def add_warning(self, msg: str):
        self.data["summary"]["warnings"].append(msg)
        logger.warning(f"  ⚠ {msg}")


# ---------------------------------------------------------------------------
# Dataclasses
# ---------------------------------------------------------------------------

@dataclass
class DocumentSection:
    """Represents a document section with hierarchy."""
    title:             str
    content:           str
    level:             int
    page_number:       Optional[int]  = None
    section_type:      str            = "text"
    section_hierarchy: Optional[List[str]] = None

    def __post_init__(self):
        if self.section_hierarchy is None:
            self.section_hierarchy = [self.title]


@dataclass
class EnrichedChunk:
    """Chunk with rich metadata."""
    text:              str
    section_title:     str
    section_hierarchy: List[str]
    page_number:       Optional[int]
    chunk_type:        str
    word_count:        int
    sentence_count:    int
    start_char:        int
    end_char:          int
    # BUG 2 FIX: headline_info is stored directly on the chunk so it survives
    # dedup, reordering, and any future filtering in main() without index drift.
    # SectionAwareChunker leaves this None — it is only populated by
    # RagasHeadlineChunker.  The Optional default keeps all existing
    # EnrichedChunk construction sites valid without changes.
    headline_info:     Optional[Dict] = None


# ---------------------------------------------------------------------------
# Helper functions
# ---------------------------------------------------------------------------

def jaccard_similarity(text_a: str, text_b: str) -> float:
    words_a = set(text_a.lower().split())
    words_b = set(text_b.lower().split())
    if not words_a or not words_b:
        return 0.0
    return len(words_a & words_b) / len(words_a | words_b)


def file_hash(path: str) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for block in iter(lambda: f.read(65536), b""):
            h.update(block)
    return h.hexdigest()


def already_indexed(client: QdrantClient, collection: str, file_hash_value: str) -> bool:
    try:
        filt = Filter(
            must=[FieldCondition(key="file_hash", match=MatchValue(value=file_hash_value))]
        )
        points, _ = client.scroll(collection_name=collection, scroll_filter=filt, limit=1)
        return len(points) > 0
    except Exception as e:
        logger.warning(f"Could not check if file is already indexed: {e}")
        return False


# ---------------------------------------------------------------------------
# BM25 tokenizer (FIX 3 — shared function used by both Ingestion and Retrieval)
# ---------------------------------------------------------------------------

# Pattern that recognises AUTOSAR compound tokens:
#   - underscore-joined identifiers:  SWS_Com_00228,  AT_231_IpduGroup
#   - double-colon namespaces:        ara::com
#   - CamelCase + digits:             ComIPduDirection  (no separator needed — kept whole)
_COMPOUND_RE = re.compile(
    r'[A-Za-z][A-Za-z0-9]*(?:[_:][A-Za-z0-9]+)+'   # underscore / colon separated
    r'|[A-Za-z]{2,}[0-9]+[A-Za-z0-9]*'              # letters directly followed by digits
    r'|[A-Za-z0-9]+(?:::[A-Za-z0-9]+)+'             # :: namespace separator
)

# Fallback: plain alphanumeric words after compound extraction
_PLAIN_RE = re.compile(r'[A-Za-z0-9]+')


def autosar_tokenize(text: str) -> List[str]:
    """
    Two-pass tokenizer that preserves AUTOSAR compound identifiers.

    Pass 1 — extract whole compound tokens and emit them lowercased.
             "SWS_Com_00228"  →  "sws_com_00228"
             "AT_231_IpduGroup" → "at_231_ipdugroup"
             "ara::com"       →  "ara::com"

    Pass 2 — split each compound into its alphanumeric parts and emit those
             too, enabling partial-identifier matching.
             "sws_com_00228"  → also emits "sws", "com", "00228"

    Plain words that are not part of any compound are emitted once as-is.

    The result is a superset of the original word_tokenize+isalnum output,
    so existing BM25 behaviour is preserved while adding compound-level hits.

    IMPORTANT: This function must be used identically in Ingestion.py and
    Evaluate_Retrieval_With_Reranker_Template.py so that query tokens match
    document tokens exactly.
    """
    tokens: List[str] = []
    covered_spans: List[Tuple[int, int]] = []

    # Pass 1 — compound tokens
    for m in _COMPOUND_RE.finditer(text):
        whole = m.group(0).lower()
        tokens.append(whole)
        covered_spans.append((m.start(), m.end()))
        # Also emit individual parts so partial queries still match
        for part in _PLAIN_RE.findall(whole):
            if len(part) > 1:           # skip single-char fragments
                tokens.append(part)

    # Pass 2 — plain alphanumeric words not already covered by a compound match
    for m in _PLAIN_RE.finditer(text):
        start, end = m.start(), m.end()
        # Skip if this span is fully inside any compound span
        if any(cs <= start and end <= ce for cs, ce in covered_spans):
            continue
        word = m.group(0).lower()
        if len(word) > 1:
            tokens.append(word)

    return tokens


# ---------------------------------------------------------------------------
# Embedding model
# ---------------------------------------------------------------------------

class OllamaBGEM3Embedder:
    """BGE-M3 embedder using Ollama with retry logic."""

    def __init__(self, base_url: str = "http://localhost:11434", model: str = "bge-m3"):
        self.base_url = base_url
        self.model    = model
        self.dimension = 1024
        self._test_connection()

    def _test_connection(self):
        try:
            response = requests.get(f"{self.base_url}/api/tags", timeout=5)
            if response.status_code == 200:
                logger.info(f"✓ Connected to Ollama at {self.base_url}")
                available_models = [m.get("name", "") for m in response.json().get("models", [])]
                if self.model not in available_models:
                    logger.warning(
                        f"Model '{self.model}' not found in Ollama. "
                        f"Available: {available_models}"
                    )
            else:
                raise ConnectionError("Ollama not responding")
        except Exception as e:
            logger.error(f"✗ Cannot connect to Ollama: {e}")
            raise

    def encode(self, texts: List[str], batch_size: int = 8,
               show_progress_bar: bool = False) -> List[Optional[List[float]]]:
        """
        Encode texts using Ollama BGE-M3 with per-text retry and automatic
        pre-truncation for oversized inputs.

        FIX 1: MAX_EMBED_CHARS is now 8000 (was 512).  BGE-M3 supports up to
        8192 tokens (~32 000 chars).  The old 512-char truncation silently made
        the second half of 90% of all chunks invisible to dense retrieval.
        """
        embeddings: List[Optional[List[float]]] = []
        skipped_count = 0
        patched_count = 0

        for i in range(0, len(texts), batch_size):
            batch = texts[i:i + batch_size]

            for text in batch:
                # Pre-truncate to MAX_EMBED_CHARS.  At 8000 chars (~2000 tokens)
                # we are well within BGE-M3's 8192-token context window.
                safe_text = text if len(text) <= MAX_EMBED_CHARS else text[:MAX_EMBED_CHARS]
                if len(text) > MAX_EMBED_CHARS:
                    logger.debug(
                        f"Chunk truncated for embedding: {len(text)} → {MAX_EMBED_CHARS} chars"
                    )

                embedding = None
                for attempt in range(3):
                    try:
                        response = requests.post(
                            f"{self.base_url}/api/embeddings",
                            json={"model": self.model, "prompt": safe_text},
                            timeout=60
                        )
                        if response.status_code == 200:
                            data = response.json()
                            if "embedding" in data:
                                raw_vec = data["embedding"]
                                if not all(math.isfinite(x) for x in raw_vec):
                                    arr = np.array(raw_vec, dtype=np.float64)
                                    non_finite_mask = ~np.isfinite(arr)
                                    n_bad = int(non_finite_mask.sum())
                                    logger.warning(
                                        f"Embedding attempt {attempt + 1}: "
                                        f"{n_bad} non-finite component(s) replaced with 0.0 "
                                        f"(chunk {len(safe_text)} chars)."
                                    )
                                    arr[non_finite_mask] = 0.0
                                    raw_vec = arr.tolist()
                                    patched_count += 1
                                embedding = raw_vec
                                break
                            else:
                                logger.warning(
                                    f"Embedding attempt {attempt + 1}: "
                                    f"unexpected response (no 'embedding' key): "
                                    f"{response.text[:120]}"
                                )
                        else:
                            body = response.text
                            logger.warning(
                                f"Embedding attempt {attempt + 1} failed "
                                f"status {response.status_code}: {body[:120]}"
                            )
                            if "unsupported value" in body:
                                logger.warning(
                                    "NaN/Inf in model output — skipping retries."
                                )
                                break
                    except Exception as e:
                        logger.warning(f"Embedding attempt {attempt + 1} error: {e}")
                    if attempt < 2:
                        time.sleep(2 ** attempt)

                if embedding is None:
                    skipped_count += 1
                    logger.warning(
                        f"Failed to embed chunk after 3 attempts "
                        f"(sent {len(safe_text)} chars, original {len(text)} chars). "
                        f"Total skipped so far: {skipped_count}"
                    )
                embeddings.append(embedding)

            if show_progress_bar and (i // batch_size) % 10 == 0:
                logger.info(f"Encoded {min(i + batch_size, len(texts))}/{len(texts)} texts")

        if patched_count:
            logger.warning(
                f"⚠  {patched_count}/{len(texts)} chunks had NaN/Inf components "
                f"replaced with 0.0."
            )
        if skipped_count:
            logger.warning(
                f"⚠  {skipped_count}/{len(texts)} chunks could not be embedded and "
                f"will be missing from the index."
            )

        return embeddings


# ---------------------------------------------------------------------------
# Chunking
# ---------------------------------------------------------------------------

class SectionAwareChunker:
    """
    Advanced chunker that respects document section structure.

    FIX 2: AUTOSAR-specific section patterns added to SECTION_PATTERNS.
    All original patterns are preserved.  The detect_sections() method
    now correctly identifies AUTOSAR test-case field labels, numbered
    sections with bracket requirement IDs, and requirement identifiers
    as section boundaries.
    """

    def __init__(self, chunk_size: int = 512, overlap: int = 256):
        self.chunk_size = chunk_size
        self.overlap    = overlap
        self.section_patterns = [re.compile(p, re.MULTILINE) for p in SECTION_PATTERNS]

    def detect_sections(self, text: str) -> List[DocumentSection]:
        """Detect document sections and hierarchy from plain text."""
        sections = []
        lines = text.split('\n')
        current_section: Dict = {"title": "Introduction", "content": "", "level": 0}
        section_stack = [current_section]

        for line in lines:
            line_stripped = line.strip()

            if not line_stripped:
                current_section["content"] += "\n"
                continue

            is_header   = False
            header_level = 0

            for pattern in self.section_patterns:
                match = pattern.match(line_stripped)
                if match:
                    is_header = True
                    if line_stripped.startswith('#'):
                        header_level = len(line_stripped) - len(line_stripped.lstrip('#'))
                    elif line_stripped.isupper():
                        header_level = 1
                    else:
                        header_level = 2
                    break

            if is_header and len(line_stripped) < 200:
                if current_section["content"].strip():
                    hierarchy = [s["title"] for s in section_stack if s["title"]]
                    section_type = (
                        "table"
                        if current_section["title"].startswith("[Table ")
                        else "text"
                    )
                    sections.append(DocumentSection(
                        title=current_section["title"],
                        content=current_section["content"].strip(),
                        level=current_section["level"],
                        section_type=section_type,
                        section_hierarchy=list(hierarchy),
                    ))

                while len(section_stack) > 1 and section_stack[-1]["level"] >= header_level:
                    section_stack.pop()

                current_section = {
                    "title": line_stripped.strip('#: ').strip(),
                    "content": "",
                    "level": header_level
                }
                section_stack.append(current_section)
            else:
                current_section["content"] += line + "\n"

        if current_section["content"].strip():
            hierarchy = [s["title"] for s in section_stack if s["title"]]
            section_type = (
                "table"
                if current_section["title"].startswith("[Table ")
                else "text"
            )
            sections.append(DocumentSection(
                title=current_section["title"],
                content=current_section["content"].strip(),
                level=current_section["level"],
                section_type=section_type,
                section_hierarchy=list(hierarchy),
            ))

        return sections if sections else [DocumentSection("Document", text, 0)]

    def _split_long_sentence(self, sentence: str) -> List[str]:
        if len(sentence) <= MAX_CHUNK_SIZE:
            return [sentence]
        words = sentence.split()
        parts: List[str] = []
        current = ""
        for word in words:
            if len(current) + len(word) + 1 > MAX_CHUNK_SIZE and current:
                parts.append(current.strip())
                current = word
            else:
                current = current + " " + word if current else word
        if current.strip():
            parts.append(current.strip())
        return parts

    def chunk_with_sentences(
        self,
        text: str,
        section_title: str = "",
        section_hierarchy: Optional[List[str]] = None
    ) -> List[EnrichedChunk]:
        """Chunk text respecting sentence boundaries with configurable overlap."""
        if not text or len(text) < MIN_CHUNK_SIZE:
            return []

        if section_hierarchy is None:
            section_hierarchy = [section_title]

        raw_sentences = sent_tokenize(text)
        sentences: List[str] = []
        for s in raw_sentences:
            sentences.extend(self._split_long_sentence(s))

        chunks: List[EnrichedChunk] = []
        current_chunk = ""
        current_start = 0

        for sentence in sentences:
            if len(current_chunk) + len(sentence) > self.chunk_size and current_chunk:
                chunks.append(EnrichedChunk(
                    text=current_chunk.strip(),
                    section_title=section_title,
                    section_hierarchy=list(section_hierarchy),
                    page_number=None,
                    chunk_type="text",
                    word_count=len(word_tokenize(current_chunk)),
                    sentence_count=len(sent_tokenize(current_chunk)),
                    start_char=current_start,
                    end_char=current_start + len(current_chunk)
                ))
                old_chunk_len = len(current_chunk)
                overlap_text = (
                    current_chunk[-self.overlap:]
                    if len(current_chunk) > self.overlap
                    else current_chunk
                )
                current_chunk = overlap_text + " " + sentence
                current_start += old_chunk_len - len(overlap_text)
            else:
                current_chunk += " " + sentence if current_chunk else sentence

        if current_chunk.strip():
            chunks.append(EnrichedChunk(
                text=current_chunk.strip(),
                section_title=section_title,
                section_hierarchy=list(section_hierarchy),
                page_number=None,
                chunk_type="text",
                word_count=len(word_tokenize(current_chunk)),
                sentence_count=len(sent_tokenize(current_chunk)),
                start_char=current_start,
                end_char=current_start + len(current_chunk)
            ))

        return chunks

    def chunk_sections(self, sections: List[DocumentSection]) -> List[EnrichedChunk]:
        """Chunk all document sections while preserving section context."""
        all_chunks: List[EnrichedChunk] = []
        for section in sections:
            hierarchy = (
                section.section_hierarchy
                if hasattr(section, 'section_hierarchy')
                else [section.title]
            )
            section_chunks = self.chunk_with_sentences(section.content, section.title, hierarchy)
            for chunk in section_chunks:
                chunk.page_number = section.page_number
                chunk.chunk_type  = section.section_type
            all_chunks.extend(section_chunks)
        return all_chunks


# ---------------------------------------------------------------------------
# RAGAS HeadlineSplitter chunker (matches generate_dataset_v2.py exactly)
# ---------------------------------------------------------------------------

def _build_httpx_clients():
    import httpx
    limits = httpx.Limits(
        max_keepalive_connections=0,
        max_connections=100,
        keepalive_expiry=0.0,
    )
    timeout = httpx.Timeout(300.0, connect=30.0)
    sync_transport  = httpx.HTTPTransport(retries=3)
    async_transport = httpx.AsyncHTTPTransport(retries=3)
    return (
        httpx.Client(limits=limits, timeout=timeout, transport=sync_transport),
        httpx.AsyncClient(limits=limits, timeout=timeout, transport=async_transport),
    )


def _build_ragas_llm(provider: str, base_url: Optional[str] = None, model: Optional[str] = None):
    from ragas.llms import LangchainLLMWrapper

    if provider == "nvidia":
        from langchain_nvidia_ai_endpoints import ChatNVIDIA
        if not os.environ.get("NVIDIA_API_KEY"):
            raise RuntimeError("NVIDIA_API_KEY not set; required for --llm-provider nvidia.")
        logger.info(f"  Headline LLM : NVIDIA NIM ({model or NVIDIA_LLM_MODEL})")
        llm = ChatNVIDIA(
            model=model or NVIDIA_LLM_MODEL,
            nvidia_api_key=os.environ["NVIDIA_API_KEY"],
            base_url=base_url or NVIDIA_BASE_URL,
            temperature=0.1,
            max_tokens=1024,
        )
    elif provider == "vllm":
        from langchain_openai import ChatOpenAI
        sync_client, async_client = _build_httpx_clients()
        logger.info(f"  Headline LLM : vLLM ({model or VLLM_LLM_MODEL}) at {base_url or VLLM_BASE_URL}")
        llm = ChatOpenAI(
            model=model or VLLM_LLM_MODEL,
            openai_api_key="dummy",
            openai_api_base=base_url or VLLM_BASE_URL,
            temperature=0.1,
            max_tokens=1024,
            http_client=sync_client,
            http_async_client=async_client,
        )
    else:
        raise ValueError(f"Unknown LLM provider: {provider!r}. Use 'nvidia' or 'vllm'.")

    return LangchainLLMWrapper(llm)


def _build_ragas_embeddings():
    from ragas.embeddings import LangchainEmbeddingsWrapper
    from langchain_huggingface import HuggingFaceEmbeddings

    logger.info(f"  RAGAS internal embed : {RAGAS_INTERNAL_EMBED_MODEL}")
    emb = HuggingFaceEmbeddings(
        model_name=RAGAS_INTERNAL_EMBED_MODEL,
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )
    return LangchainEmbeddingsWrapper(emb)


def _safe_headline_splitter():
    from ragas.testset.transforms.splitters.headline import HeadlineSplitter

    class SafeHeadlineSplitter(HeadlineSplitter):
        async def split(self, node):
            headlines = node.properties.get("headlines")
            if not headlines:
                return [], []
            return await super().split(node)

    return SafeHeadlineSplitter()


class RagasHeadlineChunker:
    """
    Alternative chunker that uses RAGAS's HeadlineSplitter — the SAME splitter
    used by generate_dataset_v2.py to produce the dataset's reference_contexts.

    USE THIS (--chunker=ragas) when ingesting documents whose evaluation dataset
    was generated by RAGAS HeadlineSplitter, so chunk boundaries in the index
    exactly match the reference_contexts the evaluator expects to find.
    """

    def __init__(self, llm, embeddings, max_workers: int = 4):
        self.llm        = llm
        self.embeddings = embeddings
        self.max_workers = max_workers

    def _run_config(self):
        from ragas.run_config import RunConfig
        return RunConfig(
            timeout=600,
            max_retries=15,
            max_wait=180,
            max_workers=self.max_workers,
            seed=42,
        )

    def chunk_file(self, pdf_path: str) -> List[EnrichedChunk]:
        """Run the full headline-extraction + split pipeline on a single PDF."""
        from langchain_community.document_loaders import PyPDFLoader
        from ragas.testset.graph import KnowledgeGraph, Node, NodeType
        from ragas.testset.transforms import apply_transforms
        from ragas.testset.transforms.extractors import HeadlinesExtractor

        loader = PyPDFLoader(pdf_path)
        docs   = loader.load()
        if not docs:
            logger.warning(f"  PyPDFLoader returned 0 pages for {pdf_path}")
            return []

        filename = os.path.basename(pdf_path)
        logger.info(f"  PyPDFLoader loaded {len(docs)} pages from {filename}")

        for d in docs:
            if "filename" not in d.metadata:
                d.metadata["filename"] = filename

        kg = KnowledgeGraph()
        for d in docs:
            kg.nodes.append(Node(
                type=NodeType.DOCUMENT,
                properties={
                    "page_content":       d.page_content,
                    "document_metadata":  d.metadata,
                },
            ))

        trans = [
            HeadlinesExtractor(llm=self.llm),
            _safe_headline_splitter(),
        ]

        logger.info(f"  Running HeadlinesExtractor + HeadlineSplitter on {len(docs)} pages...")
        try:
            apply_transforms(kg, trans, run_config=self._run_config())
        except Exception as e:
            logger.error(f"  ✗ RAGAS transforms failed for {filename}: {e}")
            return []

        # ── Audit headline extraction results ─────────────────────────────
        doc_nodes   = [n for n in kg.nodes if n.type == NodeType.DOCUMENT]
        chunk_nodes = [n for n in kg.nodes if n.type == NodeType.CHUNK]

        pages_with_headlines    = 0
        pages_without_headlines = 0
        all_extracted_headlines = []

        for node in doc_nodes:
            headlines = node.properties.get("headlines") or []
            page_meta = node.properties.get("document_metadata") or {}
            page_num  = page_meta.get("page", "?") if isinstance(page_meta, dict) else "?"
            if headlines:
                pages_with_headlines += 1
                all_extracted_headlines.extend(
                    [h if isinstance(h, str) else str(h) for h in headlines]
                )
                logger.debug(
                    f"    Page {page_num}: {len(headlines)} headline(s): "
                    f"{[str(h)[:60] for h in headlines[:3]]}"
                )
            else:
                pages_without_headlines += 1
                logger.debug(f"    Page {page_num}: NO headlines extracted")

        logger.info(
            f"  HeadlinesExtractor results: "
            f"{pages_with_headlines}/{len(doc_nodes)} pages got headlines, "
            f"{pages_without_headlines} pages got none"
        )
        logger.info(f"  Total CHUNK nodes produced: {len(chunk_nodes)}")

        if pages_without_headlines == len(doc_nodes):
            logger.warning(
                f"  ⚠ CRITICAL: HeadlinesExtractor returned NO headlines for ANY page "
                f"in {filename}. All chunks will have section_title='Document'. "
                f"Check your LLM endpoint ({self.llm}) for timeouts or errors."
            )
        elif pages_without_headlines > len(doc_nodes) * 0.5:
            logger.warning(
                f"  ⚠ HeadlinesExtractor returned no headlines for "
                f"{pages_without_headlines}/{len(doc_nodes)} pages ({filename}). "
                f"Over 50% failure rate — LLM may be overloaded or timing out."
            )

        if all_extracted_headlines:
            sample = all_extracted_headlines[:10]
            logger.info(f"  Sample headlines extracted: {sample}")

        # ── Collect CHUNK nodes ────────────────────────────────────────────
        # IMPORTANT: RAGAS HeadlineSplitter stores the extracted headlines on
        # DOCUMENT nodes, NOT on CHUNK nodes.  A CHUNK node's page_content is
        # a substring of its parent DOCUMENT node's page_content.  We recover
        # the headline for each chunk by:
        #   1. Building a list of (doc_content_prefix, headlines, page_meta)
        #      from every DOCUMENT node that has headlines.
        #   2. For each CHUNK, finding the DOCUMENT whose content contains
        #      the chunk's content as a substring — that document's first
        #      headline becomes the chunk's section_title.
        # This is necessary because node.properties.get("headlines") on a CHUNK
        # is always empty — the headlines property is never copied to child nodes.

        # Build the lookup: list of (full_doc_content_lower, headlines, page_meta)
        doc_headline_map: List[Tuple[str, List, Dict]] = []
        for dn in doc_nodes:
            dn_headlines = dn.properties.get("headlines") or []
            dn_content   = (dn.properties.get("page_content") or "").lower()
            dn_meta      = dn.properties.get("document_metadata") or {}
            if dn_headlines and dn_content:
                doc_headline_map.append((dn_content, dn_headlines, dn_meta))

        def _find_headline_for_chunk(chunk_content: str) -> Tuple[List, Dict]:
            """
            Return (headlines, page_meta) of the DOCUMENT node whose page_content
            contains this chunk's opening text.

            BUG 1 FIX (Option B — shorter probe):
            The previous 200-char probe failed for cross-page chunks because their
            content spans two PyPDFLoader DOCUMENT nodes and no single node's
            page_content contains the full 200-char prefix.  A cross-page chunk
            always *begins* on one specific page, so an 80-char probe taken from
            the start of the chunk reliably falls entirely within that page's
            content.

            Trade-off acknowledged: 80 chars is occasionally non-unique on pages
            that share the same AUTOSAR document boilerplate header.  In that case
            the first matching page wins — the headline may be from the wrong
            section within the same document, but will never be from a different
            document because the doc_headline_map only contains nodes from the
            current file's KG.
            """
            probe = chunk_content.strip().lower()[:80]
            if not probe:
                return [], {}
            for doc_content, headlines, page_meta in doc_headline_map:
                if probe in doc_content:
                    return headlines, page_meta
            return [], {}

        enriched: List[EnrichedChunk] = []
        char_cursor   = 0
        fallback_count = 0

        for node in kg.nodes:
            if node.type != NodeType.CHUNK:
                continue
            content = node.properties.get("page_content", "") or ""
            if not content.strip() or len(content) < MIN_CHUNK_SIZE:
                continue

            # Look up headline from the parent DOCUMENT node, not the CHUNK node
            found_headlines, found_meta = _find_headline_for_chunk(content)

            if found_headlines:
                first_hl      = found_headlines[0] if isinstance(found_headlines[0], str) else str(found_headlines[0])
                section_title = first_hl[:200]
                headline_info = {
                    "headlines_found": [str(h)[:80] for h in found_headlines[:5]],
                    "source_page":     found_meta.get("page"),
                }
            else:
                section_title = "Document"
                fallback_count += 1
                headline_info = {"headlines_found": [], "source_page": None}

            # Page number comes from the parent DOCUMENT metadata we already resolved
            page_num = found_meta.get("page") if found_meta else None

            # BUG 2 FIX: headline_info is attached directly to the EnrichedChunk
            # so it stays bound through dedup and any other filtering in main().
            # self._last_chunk_headline_info is no longer used.
            enriched.append(EnrichedChunk(
                text=content.strip(),
                section_title=section_title,
                section_hierarchy=[section_title],
                page_number=page_num,
                chunk_type="text",
                word_count=len(word_tokenize(content)),
                sentence_count=len(sent_tokenize(content)),
                start_char=char_cursor,
                end_char=char_cursor + len(content),
                headline_info=headline_info,
            ))
            char_cursor += len(content)

        if fallback_count:
            logger.warning(
                f"  ⚠ {fallback_count}/{len(enriched)} chunks fell back to "
                f"section_title='Document' (no headlines on their source page)."
            )

        unique_titles = set(c.section_title for c in enriched)
        logger.info(
            f"  ✓ RAGAS HeadlineSplitter produced {len(enriched)} chunks "
            f"with {len(unique_titles)} unique section titles"
        )
        if len(unique_titles) <= 1:
            logger.warning(
                f"  ⚠ Only 1 unique section title across all chunks. "
                f"Expected many distinct headline-based titles. "
                f"HeadlinesExtractor likely failed for this document."
            )
        else:
            logger.info(f"  Section title samples: {sorted(unique_titles)[:8]}")

        return enriched


# ---------------------------------------------------------------------------
# BM25 index  (FIX 3: uses autosar_tokenize instead of word_tokenize+isalnum)
# ---------------------------------------------------------------------------

class BM25Index:
    """
    BM25 sparse vector index.

    FIX 3: _tokenize() now calls autosar_tokenize() which preserves compound
    AUTOSAR identifiers (SWS_Com_00228, AT_231_IpduGroup, ara::com, etc.) as
    atomic tokens while also emitting their individual parts for partial
    matching.  This replaces the old word_tokenize+isalnum approach which
    destroyed identifier specificity by splitting on underscores and colons.

    The SAME autosar_tokenize() function is imported and used in
    Evaluate_Retrieval_With_Reranker_Template.py so that query-time tokenization
    exactly matches index-time tokenization.
    """

    def __init__(self):
        self.bm25: Optional[BM25Okapi]   = None
        self.tokenized_corpus: List[List[str]] = []
        self.vocabulary: Dict[str, int]   = {}
        self.token_idf:  Dict[str, float] = {}

    def fit(self, texts: List[str]):
        """Build BM25 index from a list of text strings."""
        self.tokenized_corpus = [self._tokenize(t) for t in texts]
        self.bm25             = BM25Okapi(self.tokenized_corpus)

        all_tokens = sorted({t for doc in self.tokenized_corpus for t in doc})
        self.vocabulary = {token: idx for idx, token in enumerate(all_tokens)}

        N = len(self.tokenized_corpus)
        for token in self.vocabulary:
            df = sum(1 for doc in self.tokenized_corpus if token in doc)
            self.token_idf[token] = np.log((N - df + 0.5) / (df + 0.5) + 1)

        logger.info(f"  ✓ BM25 vocabulary size: {len(self.vocabulary)}")

    def _tokenize(self, text: str) -> List[str]:
        """
        FIX 3: delegates to autosar_tokenize() instead of
        word_tokenize(text) + isalnum filter.
        """
        return autosar_tokenize(text)

    def get_sparse_vector(self, text: str) -> SparseVector:
        """Return a BM25 TF-IDF sparse vector for the given text."""
        tokens = self._tokenize(text)
        total  = len(tokens)
        token_counts: Dict[str, int] = {}

        for token in tokens:
            if token in self.vocabulary:
                token_counts[token] = token_counts.get(token, 0) + 1

        indices: List[int]  = []
        values:  List[float] = []

        for token, count in token_counts.items():
            tf  = count / total if total else 0.0
            idf = self.token_idf.get(token, 1.0)
            indices.append(self.vocabulary[token])
            values.append(float(tf * idf))

        return SparseVector(indices=indices, values=values)


# ---------------------------------------------------------------------------
# Document loader  (unchanged from V2)
# ---------------------------------------------------------------------------

class AdvancedDocumentLoader:
    """Enhanced document loader with pypdfium2-based PDF extraction."""

    _TOC_LINE_RE = re.compile(r'(\.\s*){2,}.*\d+\s*$')

    @staticmethod
    def _is_toc_page(page_text: str) -> bool:
        stripped = page_text.strip()
        if len(stripped) < TOC_MIN_CONTENT_CHARS:
            return True
        if len(stripped) > TOC_MAX_CONTENT_CHARS:
            return False
        non_empty_lines = [l for l in stripped.splitlines() if l.strip()]
        if len(non_empty_lines) < TOC_MIN_LINE_COUNT:
            return False
        toc_line_count = sum(
            1 for line in non_empty_lines
            if AdvancedDocumentLoader._TOC_LINE_RE.search(line)
        )
        return (toc_line_count / len(non_empty_lines)) >= TOC_LINE_RATIO

    @staticmethod
    def extract_metadata(path: str) -> Dict:
        file_stat = os.stat(path)
        return {
            "file_size_bytes":     file_stat.st_size,
            "created_timestamp":   file_stat.st_ctime,
            "modified_timestamp":  file_stat.st_mtime,
            "file_extension":      Path(path).suffix.lower(),
        }

    @staticmethod
    def load_pdf(path: str) -> Tuple[str, Dict]:
        metadata: Dict = {"num_pages": 0, "has_tables": False, "tables_count": 0}
        text_parts: List[str] = []
        skipped_pages: List[int] = []
        empty_pages:   List[int] = []
        error_pages:   List[int] = []

        pdf = pdfium.PdfDocument(path)
        try:
            num_pages = len(pdf)
            metadata["num_pages"] = num_pages
            logger.info(f"  PDF pages total: {num_pages}")

            for page_num in range(num_pages):
                page      = pdf[page_num]
                page_text = ""
                try:
                    textpage = page.get_textpage()
                    try:
                        page_text = textpage.get_text_bounded()
                    finally:
                        textpage.close()
                except Exception as page_err:
                    logger.warning(
                        f"  ⚠ Could not extract text from page {page_num + 1}: {page_err}"
                    )
                    error_pages.append(page_num + 1)
                finally:
                    page.close()

                if not page_text:
                    empty_pages.append(page_num + 1)
                    continue
                if AdvancedDocumentLoader._is_toc_page(page_text):
                    skipped_pages.append(page_num + 1)
                    continue

                text_parts.append(f"\n[Page {page_num + 1}]\n{page_text}\n")
        finally:
            pdf.close()

        if skipped_pages:
            logger.info(
                f"  ⊘ Skipped {len(skipped_pages)} TOC/boilerplate pages: "
                f"{skipped_pages[:10]}{'...' if len(skipped_pages) > 10 else ''}"
            )
        if empty_pages:
            logger.info(f"  ⊘ Empty pages (no text layer): {len(empty_pages)}")
        if error_pages:
            logger.warning(f"  ⚠ Pages with extraction errors: {error_pages}")

        content_pages = num_pages - len(skipped_pages) - len(empty_pages)
        logger.info(
            f"  PDF extraction: {content_pages}/{num_pages} pages yielded content "
            f"({len(skipped_pages)} TOC, {len(empty_pages)} empty, {len(error_pages)} errors)"
        )

        try:
            with pdfplumber.open(path) as plumber_pdf:
                for page in plumber_pdf.pages:
                    tables = page.extract_tables()
                    if tables:
                        metadata["has_tables"]   = True
                        metadata["tables_count"] += len(tables)
        except Exception as e:
            logger.warning(f"  pdfplumber table-count failed (non-fatal): {e}")

        # Attach page stats to metadata so the audit can read them
        metadata["pages_skipped_toc"]   = skipped_pages
        metadata["pages_empty"]         = empty_pages
        metadata["pages_error"]         = error_pages
        metadata["pages_with_content"]  = content_pages

        return "".join(text_parts), metadata

    @staticmethod
    def load_docx(path: str) -> Tuple[str, Dict]:
        doc        = docx.Document(path)
        text_parts: List[str] = []
        metadata: Dict = {"num_paragraphs": 0, "has_tables": False, "tables_count": 0}

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
                metadata["num_paragraphs"] += 1

        if doc.tables:
            metadata["has_tables"]   = True
            metadata["tables_count"] = len(doc.tables)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text_parts.append(row_text)

        return "\n".join(text_parts), metadata

    @staticmethod
    def load_txt(path: str) -> Tuple[str, Dict]:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        metadata = {
            "num_lines": len(text.split('\n')),
            "char_count": len(text)
        }
        return text, metadata

    @classmethod
    def load(cls, path: str) -> Tuple[Optional[str], Dict]:
        try:
            base_metadata = cls.extract_metadata(path)
            if path.lower().endswith(".pdf"):
                text, doc_metadata = cls.load_pdf(path)
            elif path.lower().endswith(".docx"):
                text, doc_metadata = cls.load_docx(path)
            elif path.lower().endswith(".txt"):
                text, doc_metadata = cls.load_txt(path)
            else:
                return None, {}
            base_metadata.update(doc_metadata)
            return text, base_metadata
        except Exception as e:
            logger.error(f"Error loading {path}: {e}")
            return None, {}


# ---------------------------------------------------------------------------
# Main ingestion pipeline
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="Advanced RAG Ingestion System V2 (Fixed)")
    parser.add_argument('--data-dir',         type=str, default=DATA_DIR)
    parser.add_argument('--collection',       type=str, default=COLLECTION)
    parser.add_argument('--qdrant-url',       type=str, default=QDRANT_URL)
    parser.add_argument('--ollama-url',       type=str, default=OLLAMA_URL)
    parser.add_argument('--ollama-model',     type=str, default=OLLAMA_MODEL)
    parser.add_argument('--chunk-size',       type=int, default=CHUNK_SIZE)
    parser.add_argument('--chunk-overlap',    type=int, default=CHUNK_OVERLAP)
    parser.add_argument('--bm25-output',      type=str, default=BM25_OUTPUT,
                        help='Path for bm25_index.json (default: ingestion_output/bm25_index.json)')
    parser.add_argument(
        '--chunker', type=str, choices=['section', 'ragas'], default='section',
        help=(
            "'section' = SectionAwareChunker (fast, regex-based). "
            "'ragas'   = RAGAS HeadlineSplitter (slow, LLM-backed). "
            "Use 'ragas' when the eval dataset was generated by RAGAS — "
            "chunk boundaries will exactly match reference_contexts."
        )
    )
    parser.add_argument('--llm-provider',     type=str, choices=['nvidia', 'vllm'], default='nvidia')
    parser.add_argument('--llm-base-url',     type=str, default=None)
    parser.add_argument('--llm-model',        type=str, default=None)
    parser.add_argument('--ragas-max-workers',type=int, default=4)
    parser.add_argument('--log-dir',          type=str, default=INGESTION_LOG_DIR,
                        help='Directory for all ingestion outputs: .log, .json audit, and bm25_index.json (default: ./ingestion_output)')
    args = parser.parse_args()

    data_dir      = args.data_dir
    collection    = args.collection
    qdrant_url    = args.qdrant_url
    ollama_url    = args.ollama_url
    ollama_model  = args.ollama_model
    chunk_size    = args.chunk_size
    chunk_overlap = args.chunk_overlap

    # ── Resolve all output paths under a single directory ─────────────────
    # If --bm25-output was explicitly supplied with a path that doesn't start
    # with the log_dir, honour it as-is.  Otherwise always write bm25_index.json
    # into the same directory as the .log and .json audit files so all ingestion
    # outputs are colocated in ./ingestion_output (or whatever --log-dir is set to).
    output_dir  = args.log_dir
    bm25_output = (
        args.bm25_output
        if os.path.isabs(args.bm25_output)
           or not args.bm25_output.startswith(("bm25_index", "./bm25_index"))
        else str(Path(output_dir) / "bm25_index.json")
    )
    # Ensure output dir exists before _setup_file_logging tries to create files in it
    Path(output_dir).mkdir(parents=True, exist_ok=True)

    # ── Wire file logging and create audit object ──────────────────────────
    log_path, json_path = _setup_file_logging(output_dir)

    run_config = {
        "timestamp":        time.strftime("%Y-%m-%d %H:%M:%S"),
        "data_dir":         data_dir,
        "collection":       collection,
        "qdrant_url":       qdrant_url,
        "ollama_url":       ollama_url,
        "ollama_model":     ollama_model,
        "chunker":          args.chunker,
        "llm_provider":     args.llm_provider,
        "llm_base_url":     args.llm_base_url,
        "llm_model":        args.llm_model,
        "chunk_size":       chunk_size,
        "chunk_overlap":    chunk_overlap,
        "max_embed_chars":  MAX_EMBED_CHARS,
        "bm25_output":      bm25_output,
        "log_file":         log_path,
        "audit_json":       json_path,
    }
    audit = IngestionAudit(json_path, run_config)

    logger.info("=" * 80)
    logger.info("ADVANCED RAG INGESTION SYSTEM V2 (FIXED)")
    logger.info("=" * 80)
    logger.info(f"MAX_EMBED_CHARS : {MAX_EMBED_CHARS}  (FIX 1 — was 512)")
    logger.info(f"Chunker         : {args.chunker}")
    logger.info(f"BM25 tokenizer  : autosar_tokenize  (FIX 3 — compound-aware)")
    logger.info(f"Output dir      : {output_dir}  (logs + audit JSON + bm25_index.json)")

    # ── Qdrant client ──────────────────────────────────────────────────────
    client = QdrantClient(url=qdrant_url)

    # ── Embedder ───────────────────────────────────────────────────────────
    if USE_OLLAMA_BGE_M3:
        try:
            logger.info(f"Using Ollama BGE-M3 model at {ollama_url}")
            embedder      = OllamaBGEM3Embedder(ollama_url, ollama_model)
            embedding_dim = embedder.dimension
        except Exception:
            logger.warning(f"Ollama unavailable, falling back to {FALLBACK_MODEL}")
            embedder      = SentenceTransformer(FALLBACK_MODEL)
            embedding_dim = 384
    else:
        logger.info(f"Using SentenceTransformer: {FALLBACK_MODEL}")
        embedder      = SentenceTransformer(FALLBACK_MODEL)
        embedding_dim = 384

    run_config["embedding_dim"]   = embedding_dim
    run_config["embedding_model"] = (
        ollama_model if isinstance(embedder, OllamaBGEM3Embedder) else FALLBACK_MODEL
    )

    # ── Collection setup ───────────────────────────────────────────────────
    existing = [c.name for c in client.get_collections().collections]
    if collection not in existing:
        client.create_collection(
            collection_name=collection,
            vectors_config={
                "dense": VectorParams(size=embedding_dim, distance=Distance.COSINE)
            },
            sparse_vectors_config={
                "bm25": SparseVectorParams(index=SparseIndexParams(on_disk=False))
            },
        )
        logger.info(f"✓ Created collection: {collection}")
    else:
        col_info     = client.get_collection(collection)
        existing_dim = None
        if hasattr(col_info.config.params.vectors, "get"):
            dense_cfg = col_info.config.params.vectors.get("dense")
            if dense_cfg is not None:
                existing_dim = dense_cfg.size
        if existing_dim is not None and existing_dim != embedding_dim:
            msg = (
                f"Dimension mismatch: collection '{collection}' has {existing_dim}-dim vectors, "
                f"but current embedder produces {embedding_dim}-dim vectors. Aborting."
            )
            logger.error(f"✗ {msg}")
            audit.add_warning(f"ABORT: {msg}")
            audit.save()
            return
        logger.info(f"✓ Collection exists: {collection}")

    # ── Chunker ────────────────────────────────────────────────────────────
    if args.chunker == "ragas":
        logger.info("\nInitializing RAGAS HeadlineSplitter chunker (LLM-backed)...")
        ragas_llm = _build_ragas_llm(
            provider=args.llm_provider,
            base_url=args.llm_base_url,
            model=args.llm_model,
        )
        chunker = RagasHeadlineChunker(
            llm=ragas_llm,
            embeddings=None,
            max_workers=args.ragas_max_workers,
        )
        run_config["ragas_max_workers"] = args.ragas_max_workers
    else:
        chunker = SectionAwareChunker(chunk_size, chunk_overlap)

    bm25_index   = BM25Index()
    total_chunks = 0
    total_files  = 0
    all_points: List[Tuple[Dict, str]] = []

    # ── Phase 1: Load, chunk, embed ────────────────────────────────────────
    logger.info("\n" + "=" * 80)
    logger.info("PHASE 1: Loading, chunking, and embedding documents")
    logger.info("=" * 80)

    all_files = []
    for root, _, files in os.walk(data_dir):
        for file in files:
            if file.lower().endswith((".pdf", ".docx", ".txt")):
                all_files.append((root, file))
    logger.info(f"Found {len(all_files)} document(s) to process in {data_dir}")

    for file_idx, (root, file) in enumerate(all_files, 1):
        path = os.path.join(root, file)
        logger.info(f"\n[{file_idx}/{len(all_files)}] → Processing: {file}")

        try:
            h = file_hash(path)
            logger.info(f"  File hash: {h[:16]}...")
            audit.start_file(file, h, path)

            if already_indexed(client, collection, h):
                logger.info("  ⊘ Already indexed — skipping")
                audit.finish_file("skipped_already_indexed")
                audit.save()
                continue

            text, doc_metadata = AdvancedDocumentLoader.load(path)

            # Record page-level stats from PDF metadata into audit
            if path.lower().endswith(".pdf"):
                audit.set_file_pages(
                    num_pages=doc_metadata.get("num_pages", 0),
                    skipped_toc=doc_metadata.get("pages_skipped_toc", []),
                )
                logger.info(
                    f"  Pages with content : {doc_metadata.get('pages_with_content', '?')} / "
                    f"{doc_metadata.get('num_pages', '?')}"
                )

            if not text or len(text.strip()) < MIN_CHUNK_SIZE:
                logger.info("  ⊘ Empty or too short after extraction — skipping")
                audit.finish_file("skipped_empty")
                audit.save()
                continue

            logger.info(f"  Extracted text length: {len(text):,} chars")

            # ── Chunking ───────────────────────────────────────────────────
            if args.chunker == "ragas":
                if not path.lower().endswith(".pdf"):
                    logger.info(f"  ⊘ --chunker=ragas only supports PDFs; skipping {file}")
                    audit.finish_file("skipped_not_pdf")
                    audit.save()
                    continue
                raw_chunks = chunker.chunk_file(path)
                logger.info(f"  ✓ RAGAS chunker produced {len(raw_chunks)} chunks")
            elif ENABLE_SECTION_AWARE:
                sections   = chunker.detect_sections(text)
                raw_chunks = chunker.chunk_sections(sections)
                unique_section_titles = set(c.section_title for c in raw_chunks)
                logger.info(
                    f"  ✓ Detected {len(sections)} sections → {len(raw_chunks)} raw chunks "
                    f"({len(unique_section_titles)} unique section titles)"
                )
                if unique_section_titles == {"Introduction"} or unique_section_titles == {"Document"}:
                    logger.warning(
                        f"  ⚠ All chunks have the same section title "
                        f"'{next(iter(unique_section_titles))}' — section detection "
                        f"likely found no matching patterns in this document."
                    )
            else:
                raw_chunks = chunker.chunk_with_sentences(text, "Document")
                logger.info(f"  ✓ Created {len(raw_chunks)} raw chunks")

            if not raw_chunks:
                logger.warning(f"  ⚠ No chunks produced for {file} — skipping")
                audit.finish_file("skipped_zero_chunks")
                audit.save()
                continue

            # ── Deduplication ──────────────────────────────────────────────
            accepted_chunks = []
            dedup_window: List[str] = []

            for chunk in raw_chunks:
                is_dup = False
                for prev_text in dedup_window[-JACCARD_WINDOW:]:
                    if jaccard_similarity(chunk.text, prev_text) > JACCARD_DEDUP_THRESHOLD:
                        is_dup = True
                        break
                if not is_dup:
                    accepted_chunks.append(chunk)
                    dedup_window.append(chunk.text)

            dedup_removed = len(raw_chunks) - len(accepted_chunks)
            if dedup_removed:
                logger.info(f"  ⊘ Removed {dedup_removed} near-duplicate chunks")

            audit.set_file_chunk_counts(
                raw=len(raw_chunks),
                dedup_removed=dedup_removed,
                final=len(accepted_chunks),
            )

            chunks = accepted_chunks
            if not chunks:
                logger.warning(f"  ⚠ All chunks removed by dedup for {file}")
                audit.finish_file("skipped_all_deduped")
                audit.save()
                continue

            # Log chunk size distribution
            chunk_sizes = [len(c.text) for c in chunks]
            logger.info(
                f"  Chunk sizes: min={min(chunk_sizes)}, "
                f"max={max(chunk_sizes)}, "
                f"avg={sum(chunk_sizes)//len(chunk_sizes)}"
            )

            # Log section title distribution
            title_counts: Dict[str, int] = {}
            for c in chunks:
                title_counts[c.section_title] = title_counts.get(c.section_title, 0) + 1
            sorted_titles = sorted(title_counts.items(), key=lambda x: -x[1])
            logger.info(
                f"  Section titles ({len(title_counts)} unique): "
                + ", ".join(f"'{t}'×{n}" for t, n in sorted_titles[:8])
                + ("..." if len(sorted_titles) > 8 else "")
            )

            chunk_texts = [chunk.text for chunk in chunks]

            # ── Embeddings ─────────────────────────────────────────────────
            logger.info(f"  ⚡ Generating embeddings for {len(chunk_texts)} chunks...")
            embed_start = time.time()

            if isinstance(embedder, OllamaBGEM3Embedder):
                embeddings = embedder.encode(chunk_texts, batch_size=8)
            else:
                raw_embeddings = embedder.encode(
                    chunk_texts,
                    batch_size=32,
                    show_progress_bar=False,
                    convert_to_numpy=True
                )
                embeddings = [emb.tolist() for emb in raw_embeddings]

            embed_elapsed = time.time() - embed_start
            embed_ok      = sum(1 for e in embeddings if e is not None)
            embed_skipped = sum(1 for e in embeddings if e is None)
            logger.info(
                f"  Embedding complete: {embed_ok} ok, {embed_skipped} failed "
                f"in {embed_elapsed:.1f}s ({embed_elapsed/max(len(chunk_texts),1)*1000:.0f}ms/chunk)"
            )

            audit.set_file_embedding_counts(skipped=embed_skipped, patched=0)

            if embed_skipped > 0:
                logger.warning(
                    f"  ⚠ {embed_skipped}/{len(chunk_texts)} chunks have no embedding "
                    f"and will be MISSING from the dense index."
                )

            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                embed_status = "ok" if embedding is not None else "skipped"

                # BM25 tokens for this chunk
                bm25_tokens = autosar_tokenize(chunk.text)

                # BUG 2 FIX: read headline_info directly from the chunk —
                # it was set on the EnrichedChunk in chunk_file() and is
                # guaranteed to be correct regardless of dedup index changes.
                audit.record_chunk(
                    chunk_id=i,
                    chunk=chunk,
                    embedding_status=embed_status,
                    bm25_tokens=bm25_tokens,
                    headline_info=chunk.headline_info,
                )

                if embedding is None:
                    logger.warning(f"  ⚠ Skipping chunk {i} — embedding failed")
                    continue

                point_id = str(uuid.uuid5(uuid.NAMESPACE_DNS, f"{h}_{i}"))
                point = {
                    "id":     point_id,
                    "vector": embedding,
                    "payload": {
                        "content":            chunk.text,
                        "source_path":        path,
                        "filename":           file,
                        "folder":             os.path.relpath(root, data_dir),
                        "file_type":          Path(file).suffix.lower(),
                        "file_hash":          h,
                        "chunk_id":           i,
                        "total_chunks":       len(chunks),
                        "chunk_length":       len(chunk.text),
                        "word_count":         chunk.word_count,
                        "sentence_count":     chunk.sentence_count,
                        "start_char":         chunk.start_char,
                        "end_char":           chunk.end_char,
                        "section_title":      chunk.section_title,
                        "section_hierarchy":  chunk.section_hierarchy,
                        "chunk_type":         chunk.chunk_type,
                        **doc_metadata
                    }
                }
                all_points.append((point, chunk.text))
                total_chunks += 1

            total_files += 1
            audit.finish_file("ok")
            # Save after every file so partial runs are recoverable
            audit.save()
            logger.info(f"  ✓ Done: {len(chunks)} chunks added (running total: {total_chunks})")

        except Exception as e:
            logger.error(f"  ✗ Error processing {file}: {e}", exc_info=True)
            if audit._current_file:
                audit.finish_file("error", error=str(e))
            audit.save()
            continue

    if not all_points:
        msg = "No documents produced any indexable chunks."
        logger.warning(msg)
        audit.add_warning(msg)
        audit.save()
        return

    # ── Phase 2: BM25 ──────────────────────────────────────────────────────
    logger.info("\n" + "=" * 80)
    logger.info(f"PHASE 2: Building BM25 index for {len(all_points)} chunks")
    logger.info("=" * 80)

    all_texts_for_bm25 = [text for _, text in all_points]
    bm25_index.fit(all_texts_for_bm25)

    audit.record_bm25(bm25_index)
    audit.save()

    bm25_data = {
        "vocabulary": bm25_index.vocabulary,
        "token_idf":  {k: float(v) for k, v in bm25_index.token_idf.items()},
    }
    with open(bm25_output, "w") as f:
        json.dump(bm25_data, f)
    logger.info(f"  ✓ Saved BM25 index to {bm25_output}")
    logger.info(f"  Vocabulary size      : {len(bm25_index.vocabulary)}")
    logger.info(f"  Compound tokens      : {len([t for t in bm25_index.vocabulary if '_' in t or '::' in t])}")

    # ── Phase 3: Upload ────────────────────────────────────────────────────
    logger.info("\n" + "=" * 80)
    logger.info("PHASE 3: Adding sparse vectors and uploading to Qdrant")
    logger.info("=" * 80)

    points_to_upload: List[PointStruct] = []
    sparse_empty_count = 0
    for point_data, chunk_text in all_points:
        sparse_vector = bm25_index.get_sparse_vector(chunk_text)
        if not sparse_vector.indices:
            sparse_empty_count += 1
        points_to_upload.append(PointStruct(
            id=point_data["id"],
            vector={
                "dense": point_data["vector"],
                "bm25":  sparse_vector
            },
            payload=point_data["payload"]
        ))

    if sparse_empty_count:
        msg = (
            f"{sparse_empty_count}/{len(points_to_upload)} chunks produced an empty "
            f"BM25 sparse vector (no vocabulary matches). These chunks will only be "
            f"retrievable via dense search."
        )
        logger.warning(f"  ⚠ {msg}")
        audit.add_warning(msg)

    batch_size  = 100
    num_batches = (len(points_to_upload) - 1) // batch_size + 1
    upload_errors = 0

    for i in range(0, len(points_to_upload), batch_size):
        batch     = points_to_upload[i:i + batch_size]
        batch_num = i // batch_size + 1
        try:
            client.upsert(collection_name=collection, points=batch, wait=True)
            logger.info(f"  ✓ Uploaded batch {batch_num}/{num_batches} ({len(batch)} points)")
            audit.record_upload_batch(batch_num, len(batch))
        except Exception as e:
            upload_errors += 1
            logger.error(f"  ✗ Upload batch {batch_num} failed: {e}")
            audit.record_upload_batch(batch_num, len(batch), error=str(e))

    if upload_errors:
        audit.add_warning(f"{upload_errors}/{num_batches} upload batches failed.")

    audit.save()

    # ── Final summary ──────────────────────────────────────────────────────
    logger.info("\n" + "=" * 80)
    logger.info("INGESTION COMPLETE")
    logger.info("=" * 80)
    logger.info(f"✓ Files processed:       {total_files}")
    logger.info(f"✓ Total chunks indexed:  {total_chunks}")
    logger.info(
        f"✓ Avg chunks/file:      "
        f"{total_chunks / total_files if total_files > 0 else 0:.1f}"
    )
    logger.info(f"✓ Collection:            {collection}")
    logger.info(f"✓ Chunker used:          {args.chunker}")
    logger.info(f"✓ Embedding dimension:   {embedding_dim}")
    logger.info(f"✓ MAX_EMBED_CHARS:       {MAX_EMBED_CHARS}  (FIX 1)")
    logger.info(f"✓ BM25 vocabulary size:  {len(bm25_index.vocabulary)}")
    logger.info(f"✓ Upload errors:         {upload_errors}")

    # Print any warnings collected
    warnings = audit.data["summary"]["warnings"]
    if warnings:
        logger.warning(f"\n{'='*80}")
        logger.warning(f"WARNINGS ({len(warnings)} total) — review these before running evaluation:")
        for w in warnings:
            logger.warning(f"  ⚠ {w}")
        logger.warning("=" * 80)
    else:
        logger.info("✓ No warnings — ingestion completed cleanly")

    files_all_doc = audit.data["summary"]["files_with_all_document_sections"]
    if files_all_doc:
        logger.warning(
            f"\n⚠ FILES WHERE ALL CHUNKS HAVE section_title='Document' ({len(files_all_doc)}):"
        )
        for f in files_all_doc:
            logger.warning(f"  - {f}")
        logger.warning(
            "  These files likely had HeadlinesExtractor failures. "
            "Re-check the LLM endpoint and retry ingestion for these files."
        )

    logger.info(f"\n✓ Full ingestion log  : {log_path}")
    logger.info(f"✓ Structured audit JSON: {json_path}")
    logger.info("=" * 80)


if __name__ == "__main__":
    main()